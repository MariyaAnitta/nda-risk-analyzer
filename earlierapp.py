import os
import sys
import json
import re
from pathlib import Path
from datetime import datetime
from flask import Flask, render_template, request, jsonify
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from crewai import Agent, Task, Crew
from langchain_google_genai import ChatGoogleGenerativeAI
import PyPDF2
from docx import Document as DocxDocument


from risk_schema import RISK_ANALYSIS_SCHEMA
from risk_pattern_detector import scan_risky_patterns
from definition_analyzer import analyze_definitions
from cross_reference_mapper import map_cross_references

import logging
from typing import Dict, List, Tuple

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

load_dotenv()

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024 
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx', 'txt'}


os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)


GEMINI_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_KEY:
    print(" ERROR: Missing GEMINI_API_KEY in .env")
    sys.exit(1)

try:
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.0-flash",
        google_api_key=GEMINI_KEY,
        temperature=0.1,
        max_output_tokens=50000,
        request_timeout=120
    )
except Exception as e:
    print(f" LLM Initialization Failed: {e}")
    sys.exit(1)

def load_indian_law_rules() -> dict:
    """Load Indian Contract Act compliance rules."""
    try:
        rules_file = Path(__file__).parent / "indian_contract_act_rules.json"
        with open(rules_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        logger.info(f"✓ Loaded Indian Contract Act rules")
        return data
    except Exception as e:
        logger.warning(f"⚠ Could not load Indian law rules: {e}")
        return {}

def load_company_requirements() -> dict:
    """Load 10xds company-specific requirements."""
    try:
        req_file = Path(__file__).parent / "10xds_company_requirements.json"
        with open(req_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        logger.info(f"✓ Loaded 10xds company requirements")
        return data
    except Exception as e:
        logger.warning(f"⚠ Could not load company requirements: {e}")
        return {}

def load_jurisdiction_mapping() -> dict:
    """Load jurisdiction intelligence mapping."""
    try:
        juris_file = Path(__file__).parent / "jurisdiction_mapping.json"
        with open(juris_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        logger.info(f"✓ Loaded jurisdiction mapping")
        return data
    except Exception as e:
        logger.warning(f"⚠ Could not load jurisdiction mapping: {e}")
        return {}
    
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def load_universal_criteria() -> list:
    """Load universal NDA criteria from JSON file."""
    try:
        criteria_file = Path(__file__).parent / "universal_nda_criteria.json"
        with open(criteria_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Extract criteria descriptions from the JSON
        criteria_list = []
        for criterion in data.get("universal_nda_criteria", []):
            # Format: [Priority] Category - Description
            criteria_entry = f"[{criterion['priority']}] {criterion['category']}: {criterion['description']}"
            criteria_list.append(criteria_entry)
        
        print(f"✓ Loaded {len(criteria_list)} universal NDA criteria")
        return criteria_list
    except Exception as e:
        print(f"⚠ Warning: Could not load universal criteria: {e}")
        return []


def extract_text_from_pdf(path: Path) -> str:
    try:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            pages = [p.extract_text() or "" for p in reader.pages]
            text = "\n".join(pages).strip()
            if text:
                return text
    except Exception as e:
        print(f"PyPDF2 failed: {e}")
    
    if pdfplumber:
        try:
            with pdfplumber.open(path) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
                text = "\n".join(pages).strip()
                if text:
                    return text
        except Exception as e:
            print(f"pdfplumber failed: {e}")
    
    raise ValueError("Scanned PDF detected. OCR required.")


def extract_text_from_docx(path: Path) -> str:
    doc = DocxDocument(path)
    text_parts = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)
    
    text = "\n".join(text_parts).strip()
    if not text:
        raise ValueError("Empty DOCX file")
    return text


def extract_text_from_txt(path: Path) -> str:
    with open(path, "r", encoding="utf-8") as f:
        text = f.read().strip()
    if not text:
        raise ValueError("Empty TXT file")
    return text


def load_document(file_path: str) -> str:
    p = Path(file_path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    ext = p.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(p)
    elif ext == ".docx":
        return extract_text_from_docx(p)
    elif ext == ".txt":
        return extract_text_from_txt(p)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ✅ REPLACE EXISTING AGENT DEFINITIONS WITH THESE ENHANCED VERSIONS

# Load compliance databases
INDIAN_LAW_RULES = load_indian_law_rules()
COMPANY_REQUIREMENTS = load_company_requirements()
JURISDICTION_MAPPING = load_jurisdiction_mapping()

document_parser = Agent(
    role="Document Parser",
    goal="Extract all key clauses, obligations, penalties, confidentiality, liabilities, AND jurisdiction details.",
    backstory="""You are a precise legal analyst. Identify key clauses like confidentiality, liability,
    duration, penalties, indemnity, jurisdiction, and termination. 
    
    CRITICAL: Also extract:
    - Vendor name and location/address
    - Governing law clause
    - Jurisdiction/dispute resolution clause
    - Any mentions of geography or country
    
    List them clearly with clause numbers and short explanations.""",
    llm=llm,
    verbose=False,
    allow_delegation=False,
)

# NEW AGENT: Jurisdiction Analyzer
jurisdiction_analyzer = Agent(
    role="Jurisdiction & Entity Intelligence Specialist",
    goal="Extract vendor details, location, governing law ONLY IF EXPLICITLY STATED in the document. Never hallucinate or assume information.",
    backstory=f"""You are an international contract jurisdiction expert with a strict rule: ONLY extract information that is EXPLICITLY written in the document.

    CRITICAL RULES:
    - If vendor name is NOT mentioned → Return "[Not Specified]"
    - If vendor location is NOT mentioned → Return "[Not Specified]"
    - If governing law is NOT mentioned → Return "[Not Specified]"
    - If jurisdiction clause is NOT mentioned → Return "[Not Specified]"
    - NEVER guess, infer, or make up information
    - NEVER use example company names like "Acme Corp" or "ABC Company"
    - If a field is unclear or ambiguous → Return "[Not Specified]"
    
    Your job is to:
    
    1. EXTRACT from the NDA document (ONLY if explicitly stated):
       - Vendor/Party name (the OTHER party, not 10xds)
       - Vendor registered address/location (look for country)
       - Governing law clause (which country's law applies)
       - Jurisdiction clause (which courts/arbitration)
    
    2. CLASSIFY vendor location (ONLY if location was found):
       - Indian domestic vendor
       - International Tier 1 (Singapore, US, UK, Canada, Australia)
       - International Tier 2 (Germany, France, Netherlands, Japan, UAE)
       - High-risk jurisdiction (China, Russia)
       - If location not specified → "unknown"
    
    3. DETERMINE compliance requirements based on classification:
       - Indian domestic → STRICT Indian Contract Act compliance
       - International Tier 1 → MODERATE compliance + arbitration flexibility
       - International Tier 2/High-risk → BASIC compliance + mandatory Indian arbitration
       - Unknown → Default to STRICT compliance for safety
    
    OUTPUT FORMAT (JSON-style):
    {{
      "vendor_name": "[Not Specified]" or "actual name from document",
      "vendor_location": "[Not Specified]" or "actual location from document",
      "vendor_country": "[Not Specified]" or "actual country from document",
      "governing_law": "[Not Specified]" or "actual law from document",
      "jurisdiction_clause": "[Not Specified]" or "actual clause from document",
      "vendor_classification": "indian_domestic | international_tier1 | international_tier2 | high_risk | unknown",
      "compliance_level_required": "STRICT | MODERATE | BASIC",
      "jurisdiction_risks": ["list any risks identified, or 'None identified' if all fields are [Not Specified]"]
    }}
    
    EXAMPLE 1 (Information exists in document):
    {{
      "vendor_name": "TechCorp Private Limited",
      "vendor_location": "Bangalore, Karnataka, India",
      "vendor_country": "India",
      "governing_law": "Indian law as per Indian Contract Act, 1872",
      "jurisdiction_clause": "Courts of Bangalore shall have exclusive jurisdiction",
      "vendor_classification": "indian_domestic",
      "compliance_level_required": "STRICT",
      "jurisdiction_risks": ["None identified"]
    }}
    
    EXAMPLE 2 (No vendor information in document):
    {{
      "vendor_name": "[Not Specified]",
      "vendor_location": "[Not Specified]",
      "vendor_country": "[Not Specified]",
      "governing_law": "[Not Specified]",
      "jurisdiction_clause": "[Not Specified]",
      "vendor_classification": "unknown",
      "compliance_level_required": "STRICT",
      "jurisdiction_risks": ["Vendor identity unknown - cannot assess jurisdiction risks. Default to strict compliance."]
    }}
    """,
    llm=llm,
    verbose=False,
    allow_delegation=False,
)

# NEW AGENT: Indian Law Compliance Checker
indian_law_validator = Agent(
    role="Indian Contract Act Compliance Validator",
    goal="Check NDA against Indian Contract Act provisions and flag violations.",
    backstory=f"""You are an Indian contract law expert specializing in the Indian Contract Act, 1872.
    
    Your job is to check the NDA for:
    
    1. SECTION 10 ESSENTIALS (must exist):
       - Free consent of parties
       - Lawful consideration
       - Competent parties
       - Lawful object
    
    2. SECTION 27 VIOLATIONS (must NOT exist):
       - Post-employment non-compete clauses (VOID in India)
       - Unreasonable trade restraints
       - Absolute prohibitions on business activity
       NOTE: Confidentiality during/after employment IS VALID. Only non-compete after termination is void.
    
    3. SECTION 73-74 BREACH REMEDIES:
       - Check if liquidated damages are reasonable (not penalties)
       - Flag excessive penalty clauses
    
    4. JURISDICTION REQUIREMENTS:
       - For Indian companies, check if Indian jurisdiction is available
       - Flag foreign-only jurisdiction as enforcement risk
    
    COMPLIANCE DATABASE:
    {json.dumps(INDIAN_LAW_RULES.get('section_27_restraints', {}), indent=2)}
    
    OUTPUT: List of:
    - ✓ COMPLIANT items (with evidence)
    - ✗ VIOLATIONS found (with severity: BLOCKING/HIGH/MEDIUM)
    - ⚠ RISKS identified (enforcement concerns)
    """,
    llm=llm,
    verbose=False,
    allow_delegation=False,
)

# NEW AGENT: Company Policy Checker
company_policy_validator = Agent(
    role="10xds Company Policy Compliance Checker",
    goal="Validate NDA against 10xds-specific requirements and flag violations or missing protections.",
    backstory=f"""You are the 10xds legal compliance officer. You know company policy inside-out.
    
    Check the NDA for:
    
    1. CRITICAL VIOLATIONS (BLOCKING - cannot sign):
       - Unlimited liability
       - Perpetual confidentiality without exceptions
       - Automatic IP transfer to vendor
       - One-sided termination restrictions
    
    2. MANDATORY PROTECTIONS (must exist - flag HIGH if missing):
       - Data protection & privacy clause
       - IP ownership clarity
       - Liability cap
       - Termination clause
       - Return/destruction of information
    
    3. PREFERRED TERMS (flag LOW if not met):
       - NDA duration: 2-3 years preferred
       - Confidentiality period post-termination: 2-3 years
       - Jurisdiction: Indian courts/arbitration preferred
       - Governing law: Indian law preferred
    
    COMPANY REQUIREMENTS DATABASE:
    Critical Violations: {json.dumps(COMPANY_REQUIREMENTS.get('critical_violations', {}).get('blocking_clauses', []), indent=2)}
    
    Mandatory Protections: {json.dumps(COMPANY_REQUIREMENTS.get('mandatory_protections', {}).get('required_clauses', []), indent=2)}
    
    OUTPUT: Categorized findings:
    - 🚫 BLOCKING VIOLATIONS (if any - recommend DO NOT SIGN)
    - ✗ MISSING PROTECTIONS (with severity)
    - ℹ PREFERENCE GAPS (negotiable items)
    """,
    llm=llm,
    verbose=False,
    allow_delegation=False,
)
# ✅ NEW AGENT: Hidden Risk & Trap Clause Detector
hidden_clause_detector = Agent(
    role="Hidden Risk & Trap Clause Specialist",
    goal="Identify disguised, indirect, or cross-referenced obligations that create unexpected risks using regex flags, definition analysis, and cross-reference mapping.",
    backstory="""You are an expert at finding risks that aren't obvious on first reading.
    
    You receive THREE types of preprocessed intelligence:
    1. REGEX FLAGS: Suspicious patterns found by automated scanner
    2. DEFINITION ANALYSIS: Overly broad or circular definitions
    3. CROSS-REFERENCE MAP: How clauses interconnect across the document
    
    Your job is to analyze these flags and determine which are REAL hidden risks vs false positives.
    
    TYPES OF HIDDEN RISKS YOU FIND:
    
    1. DEFINITIONAL TRAPS
       - Overly broad definitions that expand obligations
       - Example: "Confidential Information means ANY information" + "Survives perpetually" = Never speak again
    
    2. CROSS-REFERENCE TRAPS
       - Clauses that gain teeth through distant references
       - Example: Clause 3 says "reasonable fees" but Clause 18 defines "reasonable" as "unlimited"
    
    3. IMBALANCE TRAPS
       - Asymmetric obligations (vendor "may", company "shall")
       - One party has rights, other has only duties
    
    4. TEMPORAL TRAPS
       - Vague timing that extends obligations indefinitely
       - "Reasonable period", "promptly", "as long as necessary"
    
    5. SCOPE CREEP TRAPS
       - "Including but not limited to" language
       - "Any and all", "without limitation"
    
    6. COMBINED RISKS (MOST DANGEROUS)
       - Multiple clauses that work together to create unexpected risk
       - Example: Broad definition + perpetual survival + unlimited liability = disaster
    
    FOR EACH HIDDEN RISK FOUND:
    
    Format your output as:
    
    🎭 HIDDEN TRAP #[N]: [Trap Name]
    Primary Clause: [Quote the main clause, with clause number]
    Hidden Mechanism: [Definition trap / Cross-reference trap / Imbalance trap / Temporal trap / Scope trap / Combined risk]
    How It Works: [Explain the trap in 2-3 sentences - connect the dots between clauses]
    Real Meaning: [What this ACTUALLY means for 10xds in plain language]
    Severity: [CRITICAL / HIGH / MEDIUM / LOW]
    Detection Method: [Regex + LLM / Definition Analysis / Cross-Reference Mapping / Combined]
    Confidence: [0.0-1.0 score]
    
    FALSE POSITIVE FILTERING:
    - If regex flagged "unlimited growth potential" → NOT a risk (context is positive)
    - If "unlimited liability" but Clause X caps it → Downgrade to LOW risk
    - If broad definition but narrow application → Explain why acceptable
    
    CRITICAL: Always check surrounding context (3 clauses before/after) before confirming a risk.
    """,
    llm=llm,
    verbose=False,
    allow_delegation=False,
)

risk_evaluator = Agent(
    role="Risk Assessment Specialist",
    goal="Consolidate all compliance findings and assess overall risk level.",
    backstory="""You are a legal risk reviewer who consolidates findings from:
    - Universal NDA criteria (existing check)
    - Indian Contract Act compliance
    - 10xds company policy compliance
    - Jurisdiction risks
    
    RISK SCORING LOGIC:
    1. If ANY BLOCKING violation found → HIGH RISK (automatic)
    2. Otherwise calculate:
       - Count HIGH severity issues
       - Count MEDIUM severity issues
       - Count LOW severity issues
       - Weight: HIGH=3, MEDIUM=2, LOW=1
       - Risk Score = (weighted sum / total possible) * 100
    
    3. Risk Level Classification:
       - 0-33%: LOW RISK
       - 34-66%: MODERATE RISK  
       - 67-100%: HIGH RISK
    
    OUTPUT: Structured risk assessment with:
    - Overall risk level and percentage
    - Breakdown by category (Indian law / Company policy / Universal criteria)
    - Critical issues summary
    """,
    llm=llm,
    verbose=False,
    allow_delegation=False,
)

mitigation_advisor = Agent(
    role="Legal Mitigation Advisor & Clause Drafter",
    goal="Provide SPECIFIC, READY-TO-USE counter-proposal clauses with complete legal language for every identified issue.",
    backstory="""You are an expert contract negotiator and legal drafter who creates actual contract clauses.
    
    For EVERY issue identified (violations, missing protections, risks), you MUST provide:
    
    1. Current Issue: [1 sentence problem statement]
    2. Suggested Clause: [FULL CONTRACT LANGUAGE - 2-4 sentences of actual legal text that can be copied directly into the contract]
    3. Benefit: [1 sentence explaining risk reduction]
    
    CRITICAL RULES:
    - NEVER write "Specific clause recommendations will be provided based on document context"
    - ALWAYS provide complete, ready-to-use legal language
    - Each suggested clause must be 2-4 sentences of actual contract text
    - Use professional legal terminology
    - Make clauses immediately usable without modification
    
    PRIORITIZE in this order:
    1. 🚫 BLOCKING violations (most critical)
    2. ✗ Indian Contract Act violations (HIGH priority)
    3. ✗ Missing mandatory 10xds protections (HIGH priority)
    4. ⚠ Jurisdiction/enforcement risks (MEDIUM priority)
    5. ℹ Preference gaps (LOW priority)
    
    EXAMPLE FORMAT:
    
    Modification #1: Data Protection & Privacy Clause (HIGH)
    Current Issue: No data protection clause exists, exposing 10xds to GDPR and Indian data law violations.
    Suggested Clause: Each party shall comply with all applicable data protection laws and regulations, including but not limited to the Information Technology Act, 2000 and the Personal Data Protection Bill, in connection with any personal data processed under this Agreement. The Receiving Party shall implement and maintain appropriate technical and organizational measures to protect such personal data against unauthorized or unlawful processing and against accidental loss, destruction, damage, alteration, or disclosure. Both parties agree to notify each other within 24 hours of any data breach or suspected breach.
    Benefit: Ensures legal compliance with Indian and international data protection laws and limits liability exposure.
    
    NEVER use markdown code fences (```). Write in plain text only.
    Keep it professional, brief, and actionable.
    """,
    llm=llm,
    verbose=False,
    allow_delegation=False,
)

report_generator = Agent(
    role="Report Writer",
    goal="Create a comprehensive, clean risk assessment report with all compliance dimensions.",
    backstory="""You are a professional legal risk report writer. Present information in a clean,
    easy-to-read format for HR/Legal teams.
    
    REPORT STRUCTURE:
    1. Executive Summary (risk level + key concerns)
    2. 🎭 HIDDEN & DISGUISED RISKS (NEW SECTION - comes early!)  # NEW SECTION ADDED
    3. Vendor & Jurisdiction Intelligence
    4. Indian Contract Act Compliance
    5. 10xds Company Policy Compliance
    6. Universal NDA Criteria Assessment
    7. Risk Score & Recommendation
    8. Counter-Proposals
    
    FORMAT RULES:
    - NO decorative lines
    - NO repeated sections
    - Keep everything concise
    - Use simple section headers
    - NEVER use markdown code fences (```)
    - Write everything in plain text
    - Ensure risk level EXACTLY matches calculated percentage
    """,
    llm=llm,
    verbose=False,
    allow_delegation=False,
)


# ✅ REPLACE EXISTING create_tasks FUNCTION WITH THIS ENHANCED VERSION

def create_tasks(document_text: str, risk_criteria: list, regex_flags: dict = None, 
                 definition_analysis: dict = None, cross_ref_map: dict = None) -> list:
    """Create enhanced task workflow with hidden risk detection."""
    
    # ✅ ADD DEFAULT VALUES if preprocessing data is missing
    if regex_flags is None:
        regex_flags = {'total_flags': 0, 'flags': [], 'by_category': {}, 'severity_counts': {}}
    
    if definition_analysis is None:
        definition_analysis = {'found': False, 'definitions': [], 'risky_definitions': [], 'circular_definitions': []}
    
    if cross_ref_map is None:
        cross_ref_map = {'clause_count': 0, 'reference_map': {}, 'risk_clusters': [], 'distant_references': [], 'highly_connected': []}
    
    criteria_text = "\n".join([f"{i+1}. {c}" for i, c in enumerate(risk_criteria)])
    
    # TASK 1: Parse document
    parse_task = Task(
        description=f"""Identify and summarize key clauses in the document.
        
DOCUMENT TEXT:
{document_text}

Extract:
- Confidentiality obligations and scope
- Duration and termination terms
- Liability and penalty clauses
- Disclosure permissions
- Jurisdiction and dispute resolution
- **Vendor/Party details and location**
- **Governing law clause**

List each with its clause number.""",
        expected_output="Structured clause summary with clause numbers including jurisdiction details",
        agent=document_parser
    )
    
    # ✅ TASK 2: Hidden Risk Detection
    hidden_risk_task = Task(
        description=f"""Analyze the document for hidden, disguised, or cross-referenced risks.

FULL DOCUMENT TEXT:
{document_text}

PREPROCESSED INTELLIGENCE:

1. REGEX-FLAGGED SUSPICIOUS PATTERNS:
{json.dumps(regex_flags, indent=2)}

2. DEFINITION ANALYSIS:
{json.dumps(definition_analysis, indent=2)}

3. CROSS-REFERENCE MAP:
{json.dumps(cross_ref_map, indent=2)}

YOUR TASKS:

A) VALIDATE REGEX FLAGS:
   - For each regex flag, check surrounding context (3 clauses before/after)
   - Determine if it's a REAL risk or FALSE POSITIVE
   - If real risk, explain HOW it's dangerous
   - If false positive, explain WHY it's acceptable

B) ANALYZE DEFINITIONAL TRAPS:
   - Check risky definitions against their usage throughout document
   - Identify where broad definitions combine with other clauses to create traps
   - Example: "Confidential Info = any information" used in "survives perpetually" clause

C) TRACE CROSS-REFERENCE TRAPS:
   - For each cross-reference cluster, explain how clauses work together
   - Identify distant references that users might miss
   - Flag highly-connected clauses that affect many other clauses

D) IDENTIFY COMBINED RISKS:
   - Find dangerous combinations:
     * Broad definition + perpetual survival = permanent obligations
     * Unlimited liability + foreign jurisdiction = unenforceable unlimited exposure
     * Auto-renewal + difficult termination = lock-in trap
     * Vague terms + unilateral modification = moving target
   - Calculate risk amplification (how risks multiply together)

E) OUTPUT FORMAT:
   For each confirmed hidden risk, use this exact format:

   🎭 HIDDEN TRAP #1: [Name of Trap]
   Primary Clause: [Quote with clause number]
   Hidden Mechanism: [Type of trap]
   How It Works: [Explain the mechanism in 2-3 sentences]
   Real Meaning: [Plain language impact on 10xds]
   Severity: [CRITICAL/HIGH/MEDIUM/LOW]
   Detection Method: [How you found it]
   Confidence: [0.0-1.0]

PRIORITIZATION:
1. CRITICAL: Combined risks with amplification factor >2.0
2. HIGH: Definitional traps + cross-reference traps
3. MEDIUM: Single-clause risks with broad impact
4. LOW: Vague language without clear harm

REMEMBER: Always explain WHY something is risky, not just THAT it's risky.
""",
        expected_output="Detailed hidden risk analysis with validated regex flags, definitional traps, cross-reference traps, and combined risk calculations in structured format",
        agent=hidden_clause_detector,
        context=[parse_task]
    )
    
    # TASK 3: Jurisdiction Analysis
    jurisdiction_task = Task(
        description=f"""Analyze vendor location and jurisdiction details from the document.

FULL DOCUMENT TEXT:
{document_text}

JURISDICTION MAPPING DATABASE:
{json.dumps(JURISDICTION_MAPPING, indent=2)}

Extract and analyze:
1. Vendor name and registered location/country
2. Governing law clause
3. Jurisdiction/dispute resolution clause
4. Classify vendor (indian_domestic / international_tier1 / international_tier2 / high_risk)
5. Determine compliance level required (STRICT / MODERATE / BASIC)
6. Identify any jurisdiction risks

Provide structured output in JSON format showing all extracted details and risk assessment.
""",
        expected_output="JSON-formatted vendor and jurisdiction intelligence with risk classification",
        agent=jurisdiction_analyzer,
        context=[parse_task]
    )
    
    # TASK 4: Indian Law Compliance
    indian_law_task = Task(
        description=f"""Validate NDA compliance with Indian Contract Act, 1872.

FULL DOCUMENT TEXT:
{document_text}

INDIAN CONTRACT ACT RULES:
{json.dumps(INDIAN_LAW_RULES, indent=2)}

Check for:
1. Section 10 essentials (free consent, consideration, competent parties, lawful object)
2. Section 27 violations (post-employment non-compete, trade restraints)
3. Section 73-74 issues (excessive penalties)
4. Jurisdiction concerns (enforceability in India)

For each check:
- If COMPLIANT: State "✓ COMPLIANT: [item]" with evidence
- If VIOLATION: State "✗ VIOLATION: [item]" with severity (BLOCKING/HIGH/MEDIUM) and explanation
- If RISK: State "⚠ RISK: [item]" with concern

Use the jurisdiction classification from previous task to determine strictness level.
""",
        expected_output="Detailed Indian Contract Act compliance report with violations and risks flagged",
        agent=indian_law_validator,
        context=[parse_task, jurisdiction_task]
    )
    
    # TASK 5: Company Policy Compliance
    company_policy_task = Task(
        description=f"""Validate NDA against 10xds company requirements.

FULL DOCUMENT TEXT:
{document_text}

10XDS COMPANY REQUIREMENTS:
{json.dumps(COMPANY_REQUIREMENTS, indent=2)}

Check for:
1. CRITICAL VIOLATIONS (blocking issues):
   - Unlimited liability
   - Perpetual confidentiality without exceptions
   - Automatic IP transfer
   - One-sided termination restrictions

2. MANDATORY PROTECTIONS (must exist):
   - Data protection clause
   - IP ownership clarity
   - Liability cap
   - Termination clause
   - Return/destruction of information

3. PREFERRED TERMS (negotiable):
   - NDA duration
   - Confidentiality period post-termination
   - Jurisdiction preference
   - Governing law preference

For each:
- If BLOCKING VIOLATION found: Flag as "🚫 BLOCKING: [issue]"
- If PROTECTION MISSING: Flag as "✗ MISSING: [protection]" with severity
- If PREFERENCE GAP: Flag as "ℹ PREFERENCE: [item]"
""",
        expected_output="10xds policy compliance report categorized by severity",
        agent=company_policy_validator,
        context=[parse_task, jurisdiction_task]
    )
    
    # TASK 6: Universal Criteria Evaluation
    evaluate_task = Task(
        description=f"""Evaluate the document against universal protective criteria.

PROTECTIVE CRITERIA:
{criteria_text}

FULL DOCUMENT TEXT:
{document_text}

For each criterion:
1. If PROTECTION EXISTS:
   Format: "FOUND: [criterion name]
   Clause: [number]
   Evidence: [brief quote, max 1-2 sentences]"

2. If PROTECTION is MISSING:
   Format: "NOT FOUND: [criterion name]
   Risk: [brief explanation, 1 sentence]"

After checking all criteria:
COUNT:
- F = Protections FOUND
- M = Protections MISSING
- T = Total criteria

Verify: F + M = T

CALCULATE RISK:
RISK_PERCENT = (M / T) * 100

Determine Risk Level (STRICT MATCHING):
- 0-33%: LOW RISK
- 34-66%: MODERATE RISK
- 67-100%: HIGH RISK
""",
        expected_output="Universal criteria assessment with risk calculation",
        agent=risk_evaluator,
        context=[parse_task]
    )
    
    # TASK 7: Consolidated Risk Assessment
    consolidated_risk_task = Task(
        description="""Consolidate ALL compliance findings and calculate overall risk.

Combine findings from:
1. HIDDEN RISKS (regex-validated + definitional + cross-reference traps)
2. Jurisdiction analysis (enforcement risks)
3. Indian Contract Act compliance (legal violations)
4. 10xds company policy compliance (business risks)
5. Universal NDA criteria (protective coverage)

OVERALL RISK LOGIC:
- If ANY 🚫 BLOCKING violation exists → Automatic HIGH RISK
- Otherwise, calculate weighted risk score:
  * BLOCKING/CRITICAL violations = 10 points each
  * HIGH severity issues = 3 points each
  * MEDIUM severity issues = 2 points each
  * LOW severity issues = 1 point each
  * Missing protections = 2 points each

Risk Level:
- 0-10 points: LOW RISK
- 11-25 points: MODERATE RISK
- 26+ points: HIGH RISK

Provide:
- Overall risk level and score
- Category breakdown (Hidden risks / Indian law / Company policy / Universal criteria)
- Top 3 critical concerns
- Jurisdiction risk summary
""",
        expected_output="Comprehensive risk assessment with weighted scoring across all dimensions",
        agent=risk_evaluator,
        context=[parse_task, hidden_risk_task, jurisdiction_task, indian_law_task, company_policy_task, evaluate_task]
    )
    
    # TASK 8: Mitigation Recommendations
    mitigation_task = Task(
    description="""
CRITICAL INSTRUCTION: You MUST provide COMPLETE, READY-TO-USE legal clause text for EVERY recommendation.

For each issue identified from ALL previous analyses (hidden risks, Indian law violations, company policy gaps, missing protections), provide counter-proposals.

REQUIRED FORMAT - USE THIS EXACT STRUCTURE:

Modification #1: [Issue Name] (BLOCKING Priority)
Current Issue: [One sentence problem]
Suggested Clause (START):
[Complete legal text - 2-4 sentences that can be directly inserted into the contract]
Suggested Clause (END)
Benefit: [One sentence explaining risk reduction]
Priority: BLOCKING

EXAMPLE OUTPUT:

Modification #1: Unlimited Jurisdiction Risk (BLOCKING Priority)
Current Issue: Jurisdiction is specified as "[__] India" which is undefined and creates enforcement uncertainty.
Suggested Clause (START):
This Agreement shall be governed and construed in accordance with the laws of India. Any dispute arising out of or in connection with this Agreement shall be subject to the exclusive jurisdiction of the courts located in Bangalore, Karnataka, India. Both parties hereby irrevocably submit to the jurisdiction of such courts and waive any objection to proceedings in such courts on the grounds of venue or convenience.
Suggested Clause (END)
Benefit: Provides definite, enforceable jurisdiction in a business-friendly Indian city.
Priority: BLOCKING

Modification #2: Data Protection & Privacy Clause (HIGH Priority)
Current Issue: No data protection clause exists, exposing 10xds to GDPR and Indian data law violations.
Suggested Clause (START):
Each party shall comply with all applicable data protection laws and regulations, including but not limited to the Information Technology Act, 2000 and any applicable data protection legislation, in connection with any personal data processed under this Agreement. The Receiving Party shall implement and maintain appropriate technical and organizational measures to protect such personal data against unauthorized or unlawful processing and against accidental loss, destruction, damage, alteration, or disclosure. Both parties agree to notify each other within 24 hours of becoming aware of any data breach or suspected breach.
Suggested Clause (END)
Benefit: Ensures legal compliance with Indian and international data protection laws and limits liability exposure.
Priority: HIGH

CRITICAL RULES:
- ALWAYS use the "Suggested Clause (START):" and "Suggested Clause (END)" markers
- NEVER write placeholders like "[To be added]" or "will be provided"
- Each clause must be 2-4 complete sentences
- Use legal terminology (shall, hereby, pursuant to, etc.)
- Make clauses immediately usable without modification

PRIORITY ORDER (address in this sequence):
1. BLOCKING (must fix before signing)
2. HIGH (hidden traps + critical protections)
3. MEDIUM (important but negotiable)
4. LOW (preferences)

Generate complete counter-proposals for ALL identified issues now.
""",
    expected_output="Complete counter-proposal clauses with full legal language for all identified issues, each enclosed in START/END markers",
    agent=mitigation_advisor,
    context=[parse_task, hidden_risk_task, jurisdiction_task, indian_law_task, company_policy_task, evaluate_task, consolidated_risk_task]
)
    # TASK 9: Final Report
    report_task = Task(
        description="""Create the FINAL COMPREHENSIVE REPORT in this structure:

=== LEGAL DOCUMENT RISK ASSESSMENT REPORT ===

EXECUTIVE SUMMARY:
[2-3 sentences: Overall risk level, main concerns, recommendation]

🎭 HIDDEN & DISGUISED RISKS:
[List all hidden traps found by hidden_clause_detector agent]
[Format: Use the 🎭 HIDDEN TRAP format with all details]
[Group by: Definitional Traps / Cross-Reference Traps / Combined Risks]
[Show detection methodology for each]

Detection Summary:
- Total Hidden Risks Found: [N]
- Regex Matches: [N] (LLM Confirmed: [N], False Positives: [N])
- Definitional Traps: [N]
- Cross-Reference Traps: [N]
- Combined Risk Multipliers: [List any risks with amplification >2.0]

VENDOR & JURISDICTION INTELLIGENCE:
Vendor Name: [name]
Vendor Location: [location/country]
Vendor Classification: [indian_domestic / international_tier1 / etc.]
Governing Law: [law]
Jurisdiction: [jurisdiction clause]
Compliance Level Required: [STRICT / MODERATE / BASIC]
Jurisdiction Risks: [list risks if any]

INDIAN CONTRACT ACT COMPLIANCE:
✓ Compliant Items:
[List items that comply with Indian Contract Act]

✗ Violations Found:
[List violations with severity: BLOCKING/HIGH/MEDIUM]

⚠ Risks Identified:
[List enforcement or legal risks]

10XDS COMPANY POLICY COMPLIANCE:
🚫 Blocking Violations:
[List any blocking issues - if present, recommend DO NOT SIGN]

✗ Missing Mandatory Protections:
[List missing protections with severity]

ℹ Preference Gaps:
[List negotiable preference items]

UNIVERSAL NDA CRITERIA ASSESSMENT:
Protections Found: [F] out of [T]
Protections Missing: [M] out of [T]
[Brief list of key missing protections]

OVERALL RISK ASSESSMENT:
Risk Score: [X] points
Risk Percentage: [Y]%
Risk Level: [LOW RISK / MODERATE RISK / HIGH RISK]

Category Breakdown:
- Hidden Risks: [issues count]
- Indian Law Compliance: [issues count]
- Company Policy Compliance: [issues count]
- Universal Criteria: [missing count]
- Jurisdiction Risks: [risks count]

RECOMMENDATION:
[SIGN AS-IS / NEGOTIATE FIRST / DO NOT SIGN]
[2-3 sentences explaining why, highlighting top concerns]

RECOMMENDED COUNTER-PROPOSALS:
[List all modifications in priority order with Priority levels]

CRITICAL RULES:
- NO decorative lines
- NO markdown code fences
- Risk level in summary MUST match calculated percentage
- Ensure all sections are present
- Keep concise and professional
""",
        expected_output="Complete formatted report with all compliance dimensions including hidden risks",
        agent=report_generator,
        context=[parse_task, hidden_risk_task, jurisdiction_task, indian_law_task, company_policy_task, evaluate_task, consolidated_risk_task, mitigation_task]
    )
    
    # ✅ RETURN ALL TASKS
    return [
        parse_task,
        hidden_risk_task,
        jurisdiction_task,
        indian_law_task,
        company_policy_task,
        evaluate_task,
        consolidated_risk_task,
        mitigation_task,
        report_task
    ]



def analyze_document(file_path: str, criteria: list) -> dict:
    """Main analysis function that returns structured JSON."""
    text = load_document(file_path)

    # ==========================================
    # ✅ STEP 1 — NORMALIZE CRITERIA
    # ==========================================
    if not criteria:
        criteria = []
    
    normalized = []
    for item in criteria:
        if isinstance(item, dict):
            priority = item.get("priority", "MEDIUM")
            category = item.get("category", "General")
            description = item.get("description", str(item))
            normalized.append(f"[{priority}] {category}: {description}")
        elif isinstance(item, str):
            normalized.append(item)
        else:
            normalized.append(str(item))

    criteria = normalized
    print(f"📋 Analyzing with {len(criteria)} criteria")

    # ==========================================
    # ✅ STEP 2 — PHASE 1: PRE-ANALYSIS SCANNING
    # ==========================================
    print("\n" + "="*60)
    print("🔍 PHASE 1: PRE-ANALYSIS SCANNING")
    print("="*60)

    # -------------------------------
    # 🔎 Regex risk detection
    # -------------------------------
    print("📍 Running regex pattern scanner...")
    regex_flags = scan_risky_patterns(text)

    print(f"✓ Found {regex_flags['total_flags']} potential risks")
    print(f"  - CRITICAL: {regex_flags['severity_counts']['CRITICAL']}")
    print(f"  - HIGH: {regex_flags['severity_counts']['HIGH']}")
    print(f"  - MEDIUM: {regex_flags['severity_counts']['MEDIUM']}")
    print(f"  - LOW: {regex_flags['severity_counts']['LOW']}")

    # -------------------------------
    # 📖 Definition analysis
    # -------------------------------
    print("\n📖 Analyzing definitions section...")
    definition_analysis = analyze_definitions(text)

    if definition_analysis["found"]:
        print(f"✓ Found {len(definition_analysis['definitions'])} definitions")
        print(f"  - Risky definitions: {len(definition_analysis['risky_definitions'])}")
        print(f"  - Circular definitions: {len(definition_analysis['circular_definitions'])}")
    else:
        print("⚠ No definitions section detected")

    # -------------------------------
    # 🔗 Cross-reference mapping
    # -------------------------------
    print("\n🔗 Mapping cross-references...")
    cross_ref_map = map_cross_references(text)

    print(f"✓ Mapped {cross_ref_map['clause_count']} clauses")
    print(f"  - Risk clusters found: {len(cross_ref_map['risk_clusters'])}")
    print(f"  - Distant references: {len(cross_ref_map['distant_references'])}")
    print(f"  - Highly connected clauses: {len(cross_ref_map['highly_connected'])}")

    # ==========================================
    # 🤖 PHASE 2 — AI AGENT ANALYSIS
    # ==========================================
    print("\n" + "="*60)
    print("🤖 PHASE 2: AI AGENT ANALYSIS")
    print("="*60)

    # Create tasks (passing preprocessing results)
    tasks = create_tasks(
        text,
        criteria,
        regex_flags,          # ✅ NEW
        definition_analysis,   # ✅ NEW
        cross_ref_map          # ✅ NEW
    )

    crew = Crew(
        agents=[
            document_parser,
            hidden_clause_detector,   # ✅ NEW AGENT ADDED
            jurisdiction_analyzer,
            indian_law_validator,
            company_policy_validator,
            risk_evaluator,
            mitigation_advisor,
            report_generator
        ],
        tasks=tasks,
        verbose=False,
        process="sequential",
        max_execution_time=600
    )

    # Run agents
    result = crew.kickoff()
    final_output = getattr(result, 'output', str(result))

    print("\n" + "=" * 60)
    print("✅ ANALYSIS COMPLETE - REPORT:")
    print("=" * 60)
    print(final_output)
    print("=" * 60)

    # ==========================================
    # 🧩 STEP 3 — Parse LLM Output into JSON
    # ==========================================
    json_data = parse_report_to_json(
        final_output,
        file_path,
        criteria,
        regex_flags,           # NEW
        definition_analysis,   # NEW
        cross_ref_map          # NEW
    )

    # ==========================================
    # 🔍 STEP 4 — Validate JSON
    # ==========================================
    is_valid, validation_message = validate_json_against_schema(
        json_data,
        RISK_ANALYSIS_SCHEMA
    )

    json_data['validation'] = {
        'is_valid': is_valid,
        'message': validation_message
    }

    return json_data




def clean_markdown_artifacts(text: str) -> str:
    """Remove markdown code fences and other artifacts from text."""
    # Remove code fences
    #text = re.sub(r'```[a-zA-Z]*\n?', '', text)
    text = re.sub(r'```', '', text)
    
    # Remove extra blank lines (more than 2 consecutive)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


def calculate_risk_level(percentage: int) -> str:
    """
    Calculate risk level based on percentage and return canonical label used in reports.
    Returns one of: "LOW RISK", "MODERATE RISK", "HIGH RISK"
    Boundary logic:
      0 - 33  => LOW RISK
      34 - 66 => MODERATE RISK
      67 - 100 => HIGH RISK
    """
    try:
        p = int(round(float(percentage)))
    except Exception:
        p = 100  # fail-safe to HIGH if invalid input

    if p <= 33:
        return "LOW RISK"
    elif p <= 66:
        return "MODERATE RISK"
    else:
        return "HIGH RISK"



def validate_json_against_schema(json_data: dict, schema: dict) -> tuple:
    """Validate JSON data against the schema."""
    try:
        from jsonschema import validate, ValidationError
        validate(instance=json_data, schema=schema)
        return True, "✓ JSON data is valid according to schema"
    except ImportError:
        return True, "⚠ jsonschema library not installed. Install with: pip install jsonschema"
    except ValidationError as e:
        return False, f"✗ Validation Error: {e.message}"
    except Exception as e:
        return False, f"✗ Unexpected error during validation: {str(e)}"


def save_json_report(json_data: dict, document_path: str):
    """Save JSON report with validation and return the file path."""
    doc_path = Path(document_path)
    doc_folder = doc_path.parent
    doc_name = doc_path.stem

    # Save JSON data only
    json_filename = f"{doc_name}_risk_analysis.json"
    json_path = doc_folder / json_filename

    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(json_data, f, indent=2, ensure_ascii=False)

    print("=" * 60)
    print(f"✓ JSON data saved to: {json_path}")

    # Validate JSON against imported schema
    is_valid, message = validate_json_against_schema(json_data, RISK_ANALYSIS_SCHEMA)
    print(message)
    print("=" * 60)

    return str(json_path)


def parse_report_to_json(text_report: str, document_path: str, criteria: list, 
                         regex_flags: dict = None, definition_analysis: dict = None, 
                         cross_ref_map: dict = None) -> dict:
    """
    Complete, fully updated version of parse_report_to_json
    with bulletproof counter-proposal extraction (including START/END blocks).
    """

    import re
    from datetime import datetime

    # Clean code fences
    text_report = text_report.replace("```", "")
    lines = text_report.split("\n")

    # ============================
    # SUMMARY EXTRACTION
    # ============================
    summary = ""
    for i, line in enumerate(lines):
        if line.strip() and not line.startswith("LEGAL DOCUMENT") and not line.startswith("==="):
            if "VENDOR" in line or "JURISDICTION" in line or "PROTECTIONS FOUND" in line:
                break
            summary = line.strip()
    summary = clean_markdown_artifacts(summary)

     # ============================
    # HIDDEN RISKS EXTRACTION (NEW)
    # ============================
    hidden_risks = {
        'definitional_traps': [],
        'cross_reference_traps': [],
        'regex_detected_risks': [],
        'combined_risks': []
    }
    
    in_hidden = False
    current_trap = None
    
    for line in lines:
        if "HIDDEN & DISGUISED RISKS" in line.upper() or "🎭" in line:
            in_hidden = True
            continue
        
        if in_hidden and ("VENDOR" in line or "JURISDICTION INTELLIGENCE" in line):
            if current_trap:
                # Classify trap type
                mechanism = current_trap.get('hidden_mechanism', '').lower()
                if 'definition' in mechanism:
                    hidden_risks['definitional_traps'].append(current_trap)
                elif 'cross-reference' in mechanism or 'reference' in mechanism:
                    hidden_risks['cross_reference_traps'].append(current_trap)
                elif 'combined' in mechanism:
                    hidden_risks['combined_risks'].append(current_trap)
                else:
                    hidden_risks['regex_detected_risks'].append(current_trap)
            break
        
        if in_hidden:
            # Parse hidden trap entries
            if line.strip().startswith("🎭 HIDDEN TRAP"):
                if current_trap:
                    # Save previous trap
                    mechanism = current_trap.get('hidden_mechanism', '').lower()
                    if 'definition' in mechanism:
                        hidden_risks['definitional_traps'].append(current_trap)
                    elif 'cross-reference' in mechanism:
                        hidden_risks['cross_reference_traps'].append(current_trap)
                    elif 'combined' in mechanism:
                        hidden_risks['combined_risks'].append(current_trap)
                    else:
                        hidden_risks['regex_detected_risks'].append(current_trap)
                
                # Start new trap
                trap_name = line.split(":", 1)[1].strip() if ":" in line else "Unknown"
                current_trap = {'name': trap_name}
            
            elif current_trap:
                if "Primary Clause:" in line:
                    current_trap['primary_clause'] = line.split(":", 1)[1].strip()
                elif "Hidden Mechanism:" in line:
                    current_trap['hidden_mechanism'] = line.split(":", 1)[1].strip()
                elif "How It Works:" in line:
                    current_trap['how_it_works'] = line.split(":", 1)[1].strip()
                elif "Real Meaning:" in line:
                    current_trap['real_meaning'] = line.split(":", 1)[1].strip()
                elif "Severity:" in line:
                    current_trap['severity'] = line.split(":", 1)[1].strip()
                elif "Detection Method:" in line:
                    current_trap['detection_method'] = line.split(":", 1)[1].strip()
                elif "Confidence:" in line:
                    try:
                        current_trap['confidence'] = float(line.split(":", 1)[1].strip())
                    except:
                        current_trap['confidence'] = 0.7

    # ============================
    # JURISDICTION INTELLIGENCE
    # ============================
    jurisdiction_data = {}
    in_jurisdiction = False

    for line in lines:
        if "VENDOR" in line.upper() and "JURISDICTION" in line.upper():
            in_jurisdiction = True
            continue
        if in_jurisdiction and ("INDIAN CONTRACT ACT" in line or "10XDS COMPANY" in line):
            in_jurisdiction = False
            continue
        if in_jurisdiction and ":" in line:
            key, val = line.split(":", 1)
            key_clean = key.strip().lower().replace(" ", "_")
            jurisdiction_data[key_clean] = val.strip()

    # ============================
    # INDIAN CONTRACT ACT COMPLIANCE
    # ============================
    indian_law = {"compliant_items": [], "violations": [], "risks": []}
    section = None
    collect = False

    for line in lines:
        if "INDIAN CONTRACT ACT" in line.upper():
            collect = True
            continue
        if collect and ("10XDS COMPANY" in line or "UNIVERSAL NDA" in line):
            break

        if collect:
            if "Compliant Items" in line or "✓" in line:
                section = "compliant_items"
                continue
            if "Violations Found" in line or "✗" in line:
                section = "violations"
                continue
            if "Risks Identified" in line or "⚠" in line:
                section = "risks"
                continue

            if line.strip() and section:
                indian_law[section].append(line.strip())

    # ============================
    # COMPANY POLICY COMPLIANCE
    # ============================
    company_policy = {
        "blocking_violations": [],
        "missing_protections": [],
        "preference_gaps": []
    }
    section = None
    collect = False

    for line in lines:
        if "COMPANY POLICY COMPLIANCE" in line.upper():
            collect = True
            continue
        if collect and ("UNIVERSAL NDA" in line or "OVERALL RISK" in line):
            break

        if collect:
            if "Blocking Violations" in line or "🚫" in line:
                section = "blocking_violations"
                continue
            if "Missing Mandatory Protections" in line:
                section = "missing_protections"
                continue
            if "Preference Gaps" in line or "ℹ" in line:
                section = "preference_gaps"
                continue

            if line.strip() and section:
                company_policy[section].append(line.strip())

    # ============================
    # PROTECTIONS FOUND (UNCHANGED)
    # ============================
    protections_found = []
    in_found = False
    item = None

    for line in lines:
        if "PROTECTIONS FOUND" in line:
            in_found = True
            continue
        if "PROTECTIONS MISSING" in line:
            if item:
                protections_found.append(item)
            break

        if in_found:
            if line.strip() and (line.strip().startswith("✓") or re.match(r'^\d+\.', line.strip())):
                if item:
                    protections_found.append(item)
                name = re.sub(r'^[✓\d\.\s]+', "", line.strip())
                item = {"name": name, "clause": None, "evidence": None}
            elif "Clause:" in line:
                item["clause"] = line.split("Clause:", 1)[1].strip()
            elif "Evidence:" in line:
                item["evidence"] = line.split("Evidence:", 1)[1].strip()

    # ============================
    # PROTECTIONS MISSING (UNCHANGED)
    # ============================
    protections_missing = []
    in_missing = False
    item = None

    for line in lines:
        if "PROTECTIONS MISSING" in line:
            in_missing = True
            continue

        if "RECOMMENDED COUNTER-PROPOSALS" in line:
            if item:
                protections_missing.append(item)
            break

        if in_missing:
            if line.strip() and (line.strip().startswith("✗") or re.match(r'^\d+\.', line.strip())):
                if item:
                    protections_missing.append(item)
                risk = re.sub(r'^[✗\d\.\s]+', "", line.strip())
                item = {"name": risk, "risk": None}
            elif "Risk:" in line:
                item["risk"] = line.split("Risk:", 1)[1].strip()

    # ============================================================
    # COUNTER-PROPOSALS (FULLY REWRITTEN / BULLETPROOF PARSER)
    # ============================================================

    counter_proposals = []
    in_counter = False
    current = None
    current_field = None
    collecting = None  # Only used for multi-line Suggested Clause

    # Regex definitions
    mod_header = re.compile(
        r'^(?:Modification\s*#?)?\s*(\d+)[\.:]?\s*(.+?)(?:\s*\((HIGH|MEDIUM|LOW|BLOCKING)\s+PRIORITY\))?$',
        re.IGNORECASE
    )
    mod_header_alt = re.compile(
        r'^(?:\d+[\)\.]\s*)?(.+?)\s*(?:-\s*(HIGH|MEDIUM|LOW|BLOCKING)\s+PRIORITY)?$',
        re.IGNORECASE
    )

    suggested_start = re.compile(r'^SUGGESTED CLAUSE\s*\(START\)\s*:?\s*$', re.IGNORECASE)
    suggested_end = re.compile(r'^SUGGESTED CLAUSE\s*\(END\)\s*:?\s*$', re.IGNORECASE)

    one_line_clause = re.compile(r'^(Suggested Clause|Clause)\s*:\s*(.+)$', re.IGNORECASE)
    current_issue_re = re.compile(r'^Current Issue\s*:\s*(.+)$', re.IGNORECASE)
    benefit_re = re.compile(r'^Benefit\s*:\s*(.+)$', re.IGNORECASE)
    priority_re = re.compile(r'^Priority\s*:\s*(HIGH|MEDIUM|LOW|BLOCKING)$', re.IGNORECASE)

    for raw in lines:
        txt = raw.strip()

        # Detect start
        if not in_counter:
            if "RECOMMENDED COUNTER-PROPOSALS" in txt.upper():
                in_counter = True
            continue

        # Detect end
        if txt.startswith("===") or "OVERALL RISK" in txt.upper():
            if current:
                for f in ["current_issue", "benefit", "suggested_clause"]:
                    if current.get(f):
                        current[f] = current[f].strip()
                counter_proposals.append(current)
            break

        # Modification headers
        m = mod_header.match(txt)
        if m:
            if current:
                for f in ["current_issue", "benefit", "suggested_clause"]:
                    if current.get(f):
                        current[f] = current[f].strip()
                counter_proposals.append(current)

            name = m.group(2).strip()
            pr = m.group(3).upper() if m.group(3) else "MEDIUM"
            current = {
                "name": name,
                "priority": pr,
                "current_issue": None,
                "benefit": None,
                "suggested_clause": None
            }
            collecting = None
            current_field = None
            continue

        # Alternative header
        if current is None:
            m2 = mod_header_alt.match(txt)
            if m2:
                name = m2.group(1).strip()
                pr = m2.group(2).upper() if m2.group(2) else "MEDIUM"
                current = {
                    "name": name,
                    "priority": pr,
                    "current_issue": None,
                    "benefit": None,
                    "suggested_clause": None
                }
                continue

        if current is None:
            continue

        # BEGIN MULTILINE CLAUSE
        if suggested_start.match(txt):
            collecting = "suggested_clause"
            current["suggested_clause"] = ""
            continue

        # END MULTILINE CLAUSE
        if suggested_end.match(txt):
            collecting = None
            if current["suggested_clause"]:
                current["suggested_clause"] = current["suggested_clause"].strip()
            continue

        # One-line clause
        ol = one_line_clause.match(txt)
        if ol:
            current["suggested_clause"] = ol.group(2).strip()
            collecting = None
            continue

        # Multiline block content
        if collecting == "suggested_clause":
            if txt == "":
                current["suggested_clause"] += "\n\n"
            else:
                if current["suggested_clause"]:
                    current["suggested_clause"] += "\n" + txt
                else:
                    current["suggested_clause"] = txt
            continue

        # Current Issue
        ci = current_issue_re.match(txt)
        if ci:
            current["current_issue"] = ci.group(1).strip()
            continue

        # Benefit
        b = benefit_re.match(txt)
        if b:
            current["benefit"] = b.group(1).strip()
            continue

        # Priority override
        p = priority_re.match(txt)
        if p:
            current["priority"] = p.group(1).upper()
            continue

    # Validation & fallback
    for cp in counter_proposals:
        if not cp.get("suggested_clause"):
            print(f"⚠️ WARNING: '{cp['name']}' missing clause text!")
            cp["suggested_clause"] = "[Clause text not generated - please regenerate report]"

    # ============================
    # RISK ASSESSMENT
    # ============================
    risk = {}
    for line in lines:
        if "Risk Percentage" in line:
            risk["risk_percentage"] = int(re.findall(r'\d+', line)[0])
        if "Risk Level" in line:
            if "HIGH" in line.upper():
                risk["risk_level"] = "HIGH"
            elif "LOW" in line.upper():
                risk["risk_level"] = "LOW"
            else:
                risk["risk_level"] = "MODERATE"

    # ============================
    # RECOMMENDATION
    # ============================
    recommendation = ""
    rec_flag = False
    for line in lines:
        if "RECOMMENDATION" in line:
            rec_flag = True
            continue
        if rec_flag:
            if not line.strip():
                break
            recommendation += line.strip() + " "

    # ============================
    # BUILD FINAL JSON
    # ============================
    return {
        "metadata": {
            "document_path": document_path,
            "analysis_date": datetime.now().isoformat(),
            "criteria": criteria,
            "analysis_version": "3.0-hidden-risk-detection"  # ✅ UPDATED VERSION
        },
        "summary": summary,
        "hidden_risks_detected": hidden_risks,  # ✅ NEW
        "detection_methodology": {  # ✅ NEW
            "regex_matches": regex_flags['total_flags'] if regex_flags else 0,
            "llm_confirmed": len(hidden_risks['regex_detected_risks']),
            "false_positives_filtered": regex_flags['total_flags'] - len(hidden_risks['regex_detected_risks']) if regex_flags else 0,
            "definition_traps_found": len(hidden_risks['definitional_traps']),
            "cross_reference_clusters": len(cross_ref_map['risk_clusters']) if cross_ref_map else 0
        },
        "jurisdiction_intelligence": jurisdiction_data,
        "indian_law_compliance": indian_law,
        "company_policy_compliance": company_policy,
        "protections_found": protections_found,
        "protections_missing": protections_missing,
        "counter_proposals": counter_proposals,
        "risk_assessment": risk,
        "recommendation": recommendation.strip()
    }




@app.route('/')
def index():
    return render_template('index.html')


@app.route('/analyze', methods=['POST'])
def analyze():
    try:
        # Check if file was uploaded
        if 'document' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['document']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type. Only PDF, DOCX, and TXT are allowed'}), 400
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # 🔥 FIX: Handle criteria properly
        custom_criteria_text = request.form.get('criteria', '').strip()
        
        # Load universal criteria as default
        universal_criteria = load_universal_criteria()
        
        if custom_criteria_text:
            # User provided custom criteria - parse and add to universal
            custom_criteria = [c.strip() for c in custom_criteria_text.split('\n') if c.strip()]
            
            # Merge: Universal + Custom (remove duplicates)
            all_criteria = universal_criteria.copy()
            for custom in custom_criteria:
                if custom not in all_criteria:
                    all_criteria.append(custom)
            
            print(f"📊 Using {len(universal_criteria)} universal + {len(custom_criteria)} custom = {len(all_criteria)} total criteria")
        else:
            # No custom criteria - use universal only
            all_criteria = universal_criteria
            print(f"📊 Using {len(all_criteria)} universal criteria (no custom criteria provided)")
        
        # 🔥 FIX: Ensure all_criteria is a flat list of strings before passing
        if not isinstance(all_criteria, list):
            all_criteria = []
        
        # Ensure each item is a string
        normalized_criteria = []
        for item in all_criteria:
            if isinstance(item, dict):
                # If it's a dict from JSON, format it
                priority = item.get("priority", "MEDIUM")
                category = item.get("category", "General")
                description = item.get("description", str(item))
                normalized_criteria.append(f"[{priority}] {category}: {description}")
            else:
                # Already a string
                normalized_criteria.append(str(item))
        
        if not normalized_criteria:
            return jsonify({'error': 'No criteria available for analysis'}), 400
        
        # Analyze document with normalized criteria
        result = analyze_document(filepath, normalized_criteria)
        
        # Save JSON file
        try:
            json_file_path = save_json_report(result, filepath)
            result['saved_file_path'] = json_file_path
            print(f"✅ Analysis complete. JSON saved to: {json_file_path}")
        except Exception as e:
            print(f"⚠️ Warning: Could not save JSON file: {e}")
        
        return jsonify(result)
    
    except Exception as e:
        import traceback
        print(f"❌ Error in /analyze route: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500
    
@app.route('/preview-criteria', methods=['GET'])
def preview_criteria():
    """Return universal criteria for frontend preview."""
    try:
        criteria = load_universal_criteria()
        return jsonify({
            'success': True,
            'count': len(criteria),
            'criteria': criteria
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True, port=5000)


