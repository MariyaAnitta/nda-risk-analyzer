import os
import sys
from flask_cors import CORS
import json
import re
from pathlib import Path
from datetime import datetime
from flask import Flask, render_template, request, jsonify
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from crewai import Agent, Task, Crew
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_openai import ChatOpenAI
import PyPDF2
from docx import Document as DocxDocument
from risk_schema import RISK_ANALYSIS_SCHEMA
from risk_pattern_detector import scan_risky_patterns
from definition_analyzer import analyze_definitions
from cross_reference_mapper import map_cross_references

import logging
from typing import Dict, List, Tuple

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

#load_dotenv()
dotenv_path = Path(__file__).parent / '.env'
load_dotenv(dotenv_path=dotenv_path)


app = Flask(__name__)


app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024 
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx', 'txt'}


os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

#OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
GEMINI_KEY = os.getenv("GEMINI_API_KEY")
'''GROQ_API_KEY = os.getenv("GROQ_API_KEY")'''
if not GEMINI_KEY:
    print(" ERROR: Missing GEMINI_API_KEY in .env")
    sys.exit(1)
try:
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash-lite",
        google_api_key=GEMINI_KEY,
        temperature=0.0,
        max_output_tokens=50000,
        request_timeout=600
    )
except Exception as e:
    print(f" LLM Initialization Failed: {e}")
    sys.exit(1)

"""try:
    llm = ChatOpenAI(
        model="meta-llama/llama-3-70b-instruct",
        api_key=OPENROUTER_API_KEY,  # Correct parameter
        base_url="https://openrouter.ai/api/v1",  # Correct parameter (not openai_api_base)
        temperature=0.0,
        max_tokens=50000,
        timeout=180,
        default_headers={  #  CORRECT placement
            "HTTP-Referer": "http://localhost:5000",
            "X-Title": "NDA Risk Analyzer"
        }
    )
    print(" LLM initialized successfully with OpenRouter")
except Exception as e:
    print(f" LLM Initialization Failed: {e}")
    sys.exit(1)"""
'''try:
    llm = ChatOpenAI(
        model="llama-3.1-8b-instant",
        api_key=GROQ_API_KEY,
        base_url="https://api.groq.com/openai/v1",
        temperature=0.0,
        max_tokens=32000,  # Groq's max for this model
        timeout=180
    )
    print(" LLM initialized successfully with Groq (Llama 3.3 70B)")
    print(" Speed: 280 tokens/sec | Rate: 1K RPM | FREE")
except Exception as e:
    print(f" LLM Initialization Failed: {e}")
    sys.exit(1)'''


def load_indian_law_rules() -> dict:
    """Load Indian Contract Act compliance rules."""
    try:
        rules_file = Path(__file__).parent / "indian_contract_act_rules.json"
        with open(rules_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        logger.info(f"✓ Loaded Indian Contract Act rules")
        return data
    except Exception as e:
        logger.warning(f"⚠ Could not load Indian law rules: {e}")
        return {}

def load_company_requirements() -> dict:
    """Load 10xds company-specific requirements."""
    try:
        req_file = Path(__file__).parent / "10xds_company_requirements.json"
        with open(req_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        logger.info(f"✓ Loaded 10xds company requirements")
        return data
    except Exception as e:
        logger.warning(f"⚠ Could not load company requirements: {e}")
        return {}

def load_jurisdiction_mapping() -> dict:
    """Load jurisdiction intelligence mapping."""
    try:
        juris_file = Path(__file__).parent / "jurisdiction_mapping.json"
        with open(juris_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        logger.info(f"✓ Loaded jurisdiction mapping")
        return data
    except Exception as e:
        logger.warning(f"⚠ Could not load jurisdiction mapping: {e}")
        return {}
    
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def load_universal_criteria() -> list:
    """Load universal NDA criteria from JSON file."""
    try:
        criteria_file = Path(__file__).parent / "universal_nda_criteria.json"
        with open(criteria_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Extract criteria descriptions from the JSON
        criteria_list = []
        for criterion in data.get("universal_nda_criteria", []):
            # Format: [Priority] Category - Description
            criteria_entry = f"[{criterion['priority']}] {criterion['category']}: {criterion['description']}"
            criteria_list.append(criteria_entry)
        
        print(f"✓ Loaded {len(criteria_list)} universal NDA criteria")
        return criteria_list
    except Exception as e:
        print(f"⚠ Warning: Could not load universal criteria: {e}")
        return []


def extract_text_from_pdf(path: Path) -> str:
    try:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            pages = [p.extract_text() or "" for p in reader.pages]
            text = "\n".join(pages).strip()
            if text:
                return text
    except Exception as e:
        print(f"PyPDF2 failed: {e}")
    
    if pdfplumber:
        try:
            with pdfplumber.open(path) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
                text = "\n".join(pages).strip()
                if text:
                    return text
        except Exception as e:
            print(f"pdfplumber failed: {e}")
    
    raise ValueError("Scanned PDF detected. OCR required.")


def extract_text_from_docx(path: Path) -> str:
    doc = DocxDocument(path)
    text_parts = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)
    
    text = "\n".join(text_parts).strip()
    if not text:
        raise ValueError("Empty DOCX file")
    return text


def extract_text_from_txt(path: Path) -> str:
    with open(path, "r", encoding="utf-8") as f:
        text = f.read().strip()
    if not text:
        raise ValueError("Empty TXT file")
    return text


def load_document(file_path: str) -> str:
    p = Path(file_path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    ext = p.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(p)
    elif ext == ".docx":
        return extract_text_from_docx(p)
    elif ext == ".txt":
        return extract_text_from_txt(p)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ✅ REPLACE EXISTING AGENT DEFINITIONS WITH THESE ENHANCED VERSIONS

# Load compliance databases
INDIAN_LAW_RULES = load_indian_law_rules()
COMPANY_REQUIREMENTS = load_company_requirements()
JURISDICTION_MAPPING = load_jurisdiction_mapping()

document_analyzer = Agent(
    role="Comprehensive Document Intelligence Specialist",
    goal="Extract ALL key information: clauses, jurisdiction details, vendor info, AND hidden risks in ONE analysis pass.",
    backstory=f"""You are an expert legal analyst who performs complete document analysis in a single pass.
    
    YOUR COMPREHENSIVE ANALYSIS INCLUDES:
    
    ## SECTION 1: CLAUSE EXTRACTION
    - Identify key clauses: confidentiality, liability, duration, penalties, indemnity, jurisdiction, termination
    - Extract clause numbers and provide short explanations
    
    ## SECTION 2: JURISDICTION INTELLIGENCE
    **CRITICAL RULES - ONLY EXTRACT IF EXPLICITLY STATED:**
    - Vendor name: ONLY if mentioned → else "[Not Specified]"
    - Vendor location: ONLY if mentioned → else "[Not Specified]"
    - Governing law: ONLY if mentioned → else "[Not Specified]"
    - Jurisdiction clause: ONLY if mentioned → else "[Not Specified]"
    - NEVER guess or use example names like "Acme Corp"
    
    Classification logic:
    - If location specified → Classify as: indian_domestic | international_tier1 | international_tier2 | high_risk
    - If location NOT specified → "unknown" + default to STRICT compliance
    
    Jurisdiction mapping: {json.dumps(JURISDICTION_MAPPING, indent=2)}
    
    ## SECTION 3: HIDDEN RISK DETECTION
    You receive preprocessed intelligence:
    - REGEX FLAGS: Suspicious patterns (will be passed in task description)
    - DEFINITIONS: Overly broad definitions
    - CROSS-REFS: How clauses interconnect
    
    Identify hidden risks:
    1. DEFINITIONAL TRAPS (broad definitions + perpetual survival)
    2. CROSS-REFERENCE TRAPS (distant clause connections)
    3. IMBALANCE TRAPS (asymmetric obligations)
    4. TEMPORAL TRAPS (vague timing extending obligations)
    5. SCOPE CREEP TRAPS ("including but not limited to")
    6. COMBINED RISKS (multiple clauses creating unexpected risk)
    
    For each hidden risk found:
    
    🎭 HIDDEN TRAP #[N]: [Trap Name]
    Primary Clause: [Clause number only - NO full quote]
    Hidden Mechanism: [Type in 3-5 words]
    How It Works: [1 sentence max, 15-20 words]
    Real Meaning: [1 sentence max, 15-20 words - business impact]
    Severity: [CRITICAL/HIGH/MEDIUM/LOW]
    Detection Method: [Regex/LLM/Definition/Cross-ref]
    Confidence: [0.0-1.0]
    
    **PRIORITIZE top 5-7 risks only (skip minor LOW risks)**
    
    OUTPUT FORMAT:
    
    === SECTION 1: CLAUSE EXTRACTION ===
    [List clauses with numbers and brief explanations]
    
    === SECTION 2: JURISDICTION INTELLIGENCE ===
    {{
      "vendor_name": "[Not Specified]" or "actual name",
      "vendor_location": "[Not Specified]" or "actual location",
      "vendor_country": "[Not Specified]" or "actual country",
      "governing_law": "[Not Specified]" or "actual law",
      "jurisdiction_clause": "[Not Specified]" or "actual clause",
      "vendor_classification": "indian_domestic | international_tier1 | unknown",
      "compliance_level_required": "STRICT | MODERATE | BASIC",
      "jurisdiction_risks": ["list risks or 'None identified'"]
    }}
    
    === SECTION 3: HIDDEN RISKS ===
    [List max 7 hidden traps using compact format above]
    
    Detection Summary:
    - Total Hidden Risks Found: [N]
    - Regex Matches: [N] (LLM Confirmed: [N])
    - Definitional Traps: [N]
    - Cross-Reference Traps: [N]
    """,
    llm=llm,
    verbose=False,
    allow_delegation=False,
)

# ===================================
# AGENT 2: COMPLIANCE VALIDATOR (MEGA)
# ===================================
compliance_validator = Agent(
    role="Multi-Dimensional Compliance & Risk Assessment Specialist",
    goal="Perform ALL compliance checks (Indian law + Company policy + Universal criteria) AND calculate final risk score in ONE analysis.",
    backstory=f"""You are a comprehensive compliance expert who validates against multiple frameworks simultaneously.
    
    YOUR MULTI-DIMENSIONAL ANALYSIS:
    
    ## DIMENSION 1: INDIAN CONTRACT ACT COMPLIANCE
    
    Check for:
    1. SECTION 10 ESSENTIALS (must exist):
       - Free consent, lawful consideration, competent parties, lawful object
    
    2. SECTION 27 VIOLATIONS (must NOT exist):
       - Post-employment non-compete (VOID in India)
       - Unreasonable trade restraints
       NOTE: Confidentiality during/after employment IS VALID
    
    3. SECTION 73-74 BREACH REMEDIES:
       - Liquidated damages must be reasonable (not penalties)
    
    4. JURISDICTION REQUIREMENTS:
       - Indian companies need Indian jurisdiction available
    
    Database: {json.dumps(INDIAN_LAW_RULES.get('section_27_restraints', {}), indent=2)}
    
    OUTPUT:
    - ✅ COMPLIANT items (with evidence)
    - ❌ VIOLATIONS (severity: BLOCKING/HIGH/MEDIUM)
    - ⚠ RISKS (enforcement concerns)
    
    ## DIMENSION 2: 10XDS COMPANY POLICY COMPLIANCE
    
    Check for:
    1. CRITICAL VIOLATIONS (BLOCKING):
       - Unlimited liability
       - Perpetual confidentiality without exceptions
       - Automatic IP transfer
       - One-sided termination restrictions
    
    2. MANDATORY PROTECTIONS (HIGH if missing):
       - Data protection clause
       - IP ownership clarity
       - Liability cap
       - Termination clause
       - Return/destruction clause
    
    3. PREFERRED TERMS (LOW if not met):
       - NDA duration: 2-3 years
       - Post-termination confidentiality: 2-3 years
       - Indian courts/arbitration preferred
    
    Database: {json.dumps(COMPANY_REQUIREMENTS.get('critical_violations', {}).get('blocking_clauses', []), indent=2)}
    
    OUTPUT:
    - 🚫 BLOCKING VIOLATIONS (recommend DO NOT SIGN)
    - ❌ MISSING PROTECTIONS (with severity)
    - ℹ️ PREFERENCE GAPS (negotiable)
    
    ## DIMENSION 3: UNIVERSAL NDA CRITERIA ASSESSMENT
    
    You will receive a list of universal protective criteria.
    
    For each criterion:
    - If FOUND: "FOUND: [name] | Clause: [number] | Evidence: [1-2 sentences]"
    - If MISSING: "NOT FOUND: [name] | Risk: [1 sentence]"
    
    COUNT:
    - F = Protections FOUND
    - M = Protections MISSING
    - T = Total criteria
    Verify: F + M = T
    
    ## DIMENSION 4: OVERALL RISK CALCULATION
    
    **STRICT PERCENTAGE-BASED SCORING:**
    
    STEP 1: Calculate Weighted Points
    - BLOCKING/CRITICAL = 15 points each
    - HIGH severity = 10 points each
    - MEDIUM severity = 5 points each
    - LOW severity = 2 points each
    - Missing protections = 3 points each
    
    STEP 2: Calculate Percentage
    Risk_Percentage = (Total_Points / 100) * 100
    
    STEP 3: Determine Risk Level (STRICT BOUNDARIES)
    - 0-33%: LOW RISK
    - 34-66%: MODERATE RISK
    - 67-100%: HIGH RISK
    
    **NO OVERRIDES - Use only percentage thresholds**
    
    OUTPUT FORMAT:
    
    === INDIAN CONTRACT ACT COMPLIANCE ===
    [List compliant items, violations, risks]
    
    === 10XDS COMPANY POLICY COMPLIANCE ===
    [List blocking violations, missing protections, preference gaps]
    
    === UNIVERSAL NDA CRITERIA ===
    Protections Found: [F] out of [T]
    Protections Missing: [M] out of [T]
    [List found and missing protections]
    
    === OVERALL RISK ASSESSMENT ===
    Risk Score: [X] points
    Risk Percentage: [Y]%
    Risk Level: [LOW RISK / MODERATE RISK / HIGH RISK]
    
    Category Breakdown:
    - Hidden Risks: [count]
    - Indian Law Compliance: [count]
    - Company Policy Compliance: [count]
    - Universal Criteria: [count]
    - Jurisdiction Risks: [count]
    """,
    llm=llm,
    verbose=False,
    allow_delegation=False,
)

# ===================================
# AGENT 3: REPORT GENERATOR (MEGA)
# ===================================
report_generator = Agent(
    role="Comprehensive Report Writer & Mitigation Strategist",
    goal="Generate COMPLETE REPORT with ALL counter-proposals using ready-to-use legal clauses.",
    backstory="""You are a professional legal report writer who creates complete analysis reports with actionable mitigation strategies.
    
    YOUR COMPLETE REPORT INCLUDES:
    
    ## PART 1: COUNTER-PROPOSALS (PRIORITY SECTION)
    
    For EVERY identified issue, provide:
    
    ---PROPOSAL START---
    Name: [Short descriptive name, max 8 words]
    Priority: [BLOCKING|HIGH|MEDIUM|LOW]
    Issue: [1 sentence problem statement]
    Clause: [Complete legal text in 2-4 sentences with proper punctuation]
    Benefit: [1 sentence benefit statement]
    ---PROPOSAL END---
    
    **CRITICAL RULES:**
    - Start immediately with ---PROPOSAL START--- (NO preamble)
    - Each proposal MUST have all 5 fields
    - Clause field = 2-4 complete sentences of contract language
    - Generate 5-8 proposals minimum
    - NEVER use markdown code fences
    
    Priority order:
    1. BLOCKING/CRITICAL issues (Priority: BLOCKING)
    2. Missing mandatory protections (Priority: HIGH)
    3. HIGH severity hidden risks (Priority: HIGH)
    4. Indian law violations (Priority: HIGH/MEDIUM)
    5. Company policy gaps (Priority: MEDIUM)
    6. Preference gaps (Priority: LOW)
    
    ## PART 2: COMPLETE ANALYSIS REPORT
    
    Structure:
    
    === LEGAL DOCUMENT RISK ASSESSMENT REPORT ===
    
    EXECUTIVE SUMMARY:
    [2-3 sentences: Risk level, concerns, recommendation]
    
    🎭 HIDDEN & DISGUISED RISKS:
    [Copy hidden traps from document_analyzer output]
    
    Detection Summary:
    - Total Hidden Risks Found: [N]
    - Regex Matches: [N]
    - Definitional Traps: [N]
    - Cross-Reference Traps: [N]
    
    VENDOR & JURISDICTION INTELLIGENCE:
    [Copy jurisdiction data from document_analyzer output]
    
    INDIAN CONTRACT ACT COMPLIANCE:
    [Copy from compliance_validator output]
    
    10XDS COMPANY POLICY COMPLIANCE:
    [Copy from compliance_validator output]
    
    UNIVERSAL NDA CRITERIA ASSESSMENT:
    [Copy from compliance_validator output]
    
    OVERALL RISK ASSESSMENT:
    [Copy risk calculation from compliance_validator output]
    
    RECOMMENDATION:
    [SIGN AS-IS / NEGOTIATE FIRST / DO NOT SIGN]
    [2-3 sentences explaining why]
    
    RECOMMENDED COUNTER-PROPOSALS:
    [CRITICAL: Copy ALL proposals from Part 1 with EXACT ---PROPOSAL START/END--- format]
    [Minimum 5 proposals required]
    
    **FORMAT RULES:**
    - NO decorative lines
    - NO markdown code fences
    - Risk level must EXACTLY match percentage
    - All sections present
    - Concise and professional
    """,
    llm=llm,
    verbose=False,
    allow_delegation=False,
)

def create_tasks(document_text: str, risk_criteria: list, regex_flags: dict = None, 
                 definition_analysis: dict = None, cross_ref_map: dict = None) -> list:
    """
    TOKEN-OPTIMIZED: Reduces token usage by 50% while maintaining analysis quality.
    
    KEY OPTIMIZATION: Only Task 1 receives full document. Task 2 and 3 work from context.
    """
    
    # Provide defaults
    if regex_flags is None:
        regex_flags = {'total_flags': 0, 'flags': [], 'by_category': {}, 'severity_counts': {}}
    
    if definition_analysis is None:
        definition_analysis = {'found': False, 'definitions': [], 'risky_definitions': [], 'circular_definitions': []}
    
    if cross_ref_map is None:
        cross_ref_map = {'clause_count': 0, 'reference_map': {}, 'risk_clusters': [], 'distant_references': [], 'highly_connected': []}
    
    criteria_text = "\n".join([f"{i+1}. {c}" for i, c in enumerate(risk_criteria)])
    
    # ==========================================
    # TASK 1: DOCUMENT ANALYSIS (MEGA TASK)
    # ==========================================
    document_analysis_task = Task(
        description=f"""Perform COMPLETE document analysis in ONE pass.

FULL DOCUMENT TEXT:
{document_text}

PREPROCESSED INTELLIGENCE:
1. REGEX FLAGS: {json.dumps(regex_flags, indent=2)}
2. DEFINITIONS: {json.dumps(definition_analysis, indent=2)}
3. CROSS-REFS: {json.dumps(cross_ref_map, indent=2)}

JURISDICTION MAPPING: {json.dumps(JURISDICTION_MAPPING, indent=2)}

YOUR ANALYSIS MUST INCLUDE ALL 3 SECTIONS:

=== SECTION 1: CLAUSE EXTRACTION ===
Extract and summarize key clauses with clause numbers:
- Confidentiality obligations and scope
- Duration and termination terms
- Liability and penalty clauses
- Disclosure permissions
- Jurisdiction and dispute resolution
- Vendor/Party details and location
- Governing law clause

=== SECTION 2: JURISDICTION INTELLIGENCE ===
Extract ONLY if explicitly stated (use "[Not Specified]" if not found):
- Vendor name
- Vendor location/country
- Governing law clause
- Jurisdiction clause
- Vendor classification
- Compliance level required
- Jurisdiction risks

=== SECTION 3: HIDDEN RISKS ===
Analyze preprocessed intelligence for hidden traps.
Use COMPACT format (max 7 risks):
- Primary Clause: [Clause number only]
- Hidden Mechanism: [Type in 3-5 words]
- How It Works: [1 sentence, 15-20 words]
- Real Meaning: [1 sentence, 15-20 words]
- Severity: [CRITICAL/HIGH/MEDIUM/LOW]

Provide complete output with all 3 sections clearly separated.
""",
        expected_output="Complete document intelligence with clauses, jurisdiction data, and hidden risks",
        agent=document_analyzer
    )
    
    # ==========================================
    # TASK 2: COMPLIANCE VALIDATION (MEGA TASK)
    # 🔑 KEY OPTIMIZATION: NO DOCUMENT TEXT - Uses context from Task 1
    # ==========================================
    compliance_validation_task = Task(
        description=f"""Perform COMPLETE compliance analysis against ALL frameworks in ONE pass.

⚠️ IMPORTANT: You will receive the COMPLETE DOCUMENT ANALYSIS from the previous task via context.
This includes:
- All extracted clauses with clause numbers
- Jurisdiction intelligence
- Hidden risks identified
- Full clause details

DO NOT request the document text - use the clause information provided in context.

UNIVERSAL PROTECTIVE CRITERIA:
{criteria_text}

INDIAN CONTRACT ACT RULES:
{json.dumps(INDIAN_LAW_RULES, indent=2)}

10XDS COMPANY REQUIREMENTS:
{json.dumps(COMPANY_REQUIREMENTS, indent=2)}

YOUR ANALYSIS MUST INCLUDE ALL 4 DIMENSIONS:

=== DIMENSION 1: INDIAN CONTRACT ACT COMPLIANCE ===
Using the clauses from context, check:
- Section 10 essentials (free consent, consideration, competent parties, lawful object)
- Section 27 violations (post-employment non-compete, trade restraints)
- Section 73-74 (liquidated damages must be reasonable)
- Jurisdiction requirements (Indian companies need Indian jurisdiction)

Output:
- ✅ COMPLIANT items (with evidence from clauses provided)
- ❌ VIOLATIONS (severity: BLOCKING/HIGH/MEDIUM)
- ⚠️ RISKS (enforcement concerns)

=== DIMENSION 2: 10XDS COMPANY POLICY COMPLIANCE ===
Using the clauses from context, check:
1. CRITICAL VIOLATIONS (BLOCKING):
   - Unlimited liability
   - Perpetual confidentiality without exceptions
   - Automatic IP transfer
   - One-sided termination restrictions

2. MANDATORY PROTECTIONS (HIGH if missing):
   - Data protection clause
   - IP ownership clarity
   - Liability cap
   - Termination clause
   - Return/destruction clause

3. PREFERRED TERMS (LOW if not met):
   - NDA duration: 2-3 years
   - Post-termination confidentiality: 2-3 years
   - Indian courts/arbitration preferred

Output:
- 🚫 BLOCKING VIOLATIONS (recommend DO NOT SIGN)
- ❌ MISSING PROTECTIONS (with severity)
- ℹ️ PREFERENCE GAPS (negotiable)

=== DIMENSION 3: UNIVERSAL NDA CRITERIA ASSESSMENT ===
For each criterion in the list above, check against clauses from context:
- If FOUND: "FOUND: [name] | Clause: [number from context] | Evidence: [brief quote]"
- If MISSING: "NOT FOUND: [name] | Risk: [explanation]"

Count and verify: F + M = T

=== DIMENSION 4: OVERALL RISK CALCULATION ===
Calculate risk using weighted points:
- BLOCKING/CRITICAL = 15 points
- HIGH = 10 points
- MEDIUM = 5 points
- LOW = 2 points
- Missing protection = 3 points

Risk Percentage = (Total Points / 100) * 100

Determine Risk Level (STRICT):
- 0-33%: LOW RISK
- 34-66%: MODERATE RISK
- 67-100%: HIGH RISK

Provide complete output with all 4 dimensions clearly separated, including:
- Point breakdown showing calculation
- Risk percentage
- Risk level matching percentage bracket
- Category breakdown (Hidden risks, Indian law, Company policy, Universal criteria, Jurisdiction)
""",
        expected_output="Complete multi-dimensional compliance analysis with risk calculation",
        agent=compliance_validator,
        context=[document_analysis_task]  # ← This passes Task 1 output automatically
    )
    
    # ==========================================
    # TASK 3: REPORT GENERATION (MEGA TASK)
    # ==========================================
    report_generation_task = Task(
        description=f"""Generate COMPLETE FINAL REPORT with counter-proposals and analysis.

⚠️ IMPORTANT: You have access to:
- Task 1 output: Complete document analysis (clauses, jurisdiction, hidden risks)
- Task 2 output: Complete compliance validation (all 4 dimensions + risk score)

Use this context to build the report. DO NOT request document text.

CRITICAL: Start IMMEDIATELY with counter-proposals (no preamble).

## PART 1: COUNTER-PROPOSALS (START HERE)

Generate counter-proposals for ALL identified issues using EXACT format:

---PROPOSAL START---
Name: [Short name, max 8 words]
Priority: [BLOCKING|HIGH|MEDIUM|LOW]
Issue: [1 sentence problem statement]
Clause: [2-4 sentences of legal text]
Benefit: [1 sentence benefit statement]
---PROPOSAL END---

Generate 5-8 proposals minimum in priority order:
1. BLOCKING/CRITICAL issues
2. Missing mandatory protections
3. HIGH severity hidden risks
4. Indian law violations
5. Company policy gaps
6. Preference gaps

## PART 2: COMPLETE ANALYSIS REPORT

After proposals, generate structured report:

=== LEGAL DOCUMENT RISK ASSESSMENT REPORT ===

EXECUTIVE SUMMARY:
[2-3 sentences with risk level and recommendation]

🎭 HIDDEN & DISGUISED RISKS:
[Copy from document_analysis_task context]

Detection Summary:
- Total Hidden Risks Found: [N]
- Regex Matches: [N]
- Definitional Traps: [N]
- Cross-Reference Traps: [N]

VENDOR & JURISDICTION INTELLIGENCE:
[Copy from document_analysis_task context]

INDIAN CONTRACT ACT COMPLIANCE:
[Copy from compliance_validation_task context]

10XDS COMPANY POLICY COMPLIANCE:
[Copy from compliance_validation_task context]

UNIVERSAL NDA CRITERIA ASSESSMENT:
[Copy from compliance_validation_task context]

OVERALL RISK ASSESSMENT:
[Copy risk calculation from compliance_validation_task context]

RECOMMENDATION:
[SIGN AS-IS / NEGOTIATE FIRST / DO NOT SIGN]
[2-3 sentences explaining why]

RECOMMENDED COUNTER-PROPOSALS:
[CRITICAL: Copy ALL proposals from Part 1 with EXACT ---PROPOSAL START/END--- format]
[Minimum 5 proposals required]

**FORMAT RULES:**
- NO markdown code fences
- Risk level must EXACTLY match percentage
- All sections present
- Concise and professional
""",
        expected_output="Complete formatted report with counter-proposals in ---PROPOSAL START/END--- format and all analysis sections",
        agent=report_generator,
        context=[document_analysis_task, compliance_validation_task]  # ← Access to both previous outputs
    )
    
    return [
        document_analysis_task,
        compliance_validation_task,
        report_generation_task
    ]




def analyze_document(file_path: str, criteria: list) -> dict:
    """Main analysis function that returns structured JSON."""
    text = load_document(file_path)

    # ==========================================
    # ✅ STEP 1 — NORMALIZE CRITERIA
    # ==========================================
    if not criteria:
        criteria = []
    
    normalized = []
    for item in criteria:
        if isinstance(item, dict):
            priority = item.get("priority", "MEDIUM")
            category = item.get("category", "General")
            description = item.get("description", str(item))
            normalized.append(f"[{priority}] {category}: {description}")
        elif isinstance(item, str):
            normalized.append(item)
        else:
            normalized.append(str(item))

    criteria = normalized
    print(f"📋 Analyzing with {len(criteria)} criteria")

    # ==========================================
    # ✅ STEP 2 — PHASE 1: PRE-ANALYSIS SCANNING
    # ==========================================
    print("\n" + "="*60)
    print("🔍 PHASE 1: PRE-ANALYSIS SCANNING")
    print("="*60)

    # -------------------------------
    # 🔎 Regex risk detection
    # -------------------------------
    print("📍 Running regex pattern scanner...")
    regex_flags = scan_risky_patterns(text)

    print(f"✓ Found {regex_flags['total_flags']} potential risks")
    print(f"  - CRITICAL: {regex_flags['severity_counts']['CRITICAL']}")
    print(f"  - HIGH: {regex_flags['severity_counts']['HIGH']}")
    print(f"  - MEDIUM: {regex_flags['severity_counts']['MEDIUM']}")
    print(f"  - LOW: {regex_flags['severity_counts']['LOW']}")

    # -------------------------------
    # 📖 Definition analysis
    # -------------------------------
    print("\n📖 Analyzing definitions section...")
    definition_analysis = analyze_definitions(text)

    if definition_analysis["found"]:
        print(f"✓ Found {len(definition_analysis['definitions'])} definitions")
        print(f"  - Risky definitions: {len(definition_analysis['risky_definitions'])}")
        print(f"  - Circular definitions: {len(definition_analysis['circular_definitions'])}")
    else:
        print("⚠ No definitions section detected")

    # -------------------------------
    # 🔗 Cross-reference mapping
    # -------------------------------
    print("\n🔗 Mapping cross-references...")
    cross_ref_map = map_cross_references(text)

    print(f"✓ Mapped {cross_ref_map['clause_count']} clauses")
    print(f"  - Risk clusters found: {len(cross_ref_map['risk_clusters'])}")
    print(f"  - Distant references: {len(cross_ref_map['distant_references'])}")
    print(f"  - Highly connected clauses: {len(cross_ref_map['highly_connected'])}")

    # ==========================================
    # 🤖 PHASE 2 — AI AGENT ANALYSIS
    # ==========================================
    print("\n" + "="*60)
    print("🤖 PHASE 2: AI AGENT ANALYSIS")
    print("="*60)

    # Create tasks (passing preprocessing results)
    tasks = create_tasks(
        text,
        criteria,
        regex_flags,          # ✅ NEW
        definition_analysis,   # ✅ NEW
        cross_ref_map          # ✅ NEW
    )

    crew = Crew(
    agents=[
        document_analyzer,      # Agent 1: Mega analysis agent
        compliance_validator,   # Agent 2: Mega compliance agent
        report_generator        # Agent 3: Mega report agent
    ],
    tasks=tasks,  # Now only 3 tasks instead of 9
    verbose=False,
    process="sequential",
    max_execution_time=600
)

    # Run agents
    result = crew.kickoff()
    final_output = getattr(result, 'output', str(result))
    debug_file = Path("debug_llm_output.txt")
    with open(debug_file, 'w', encoding='utf-8') as f:
        f.write("=== FULL LLM OUTPUT ===\n\n")
        f.write(final_output)
        f.write("\n\n=== END OUTPUT ===")
    logger.info(f"📝 Saved raw LLM output to {debug_file}")

    # Check if counter-proposals section exists
    if "RECOMMENDED COUNTER-PROPOSALS" in final_output:
        # Extract just the counter-proposals section for analysis
        proposals_start = final_output.find("RECOMMENDED COUNTER-PROPOSALS")
        proposals_section = final_output[proposals_start:proposals_start+2000]
        logger.info(f"📋 Counter-proposals section preview:\n{proposals_section}")
    else:
        logger.error("❌ No counter-proposals section found in LLM output!")
        logger.error(f"Output length: {len(final_output)} chars")
        logger.error(f"Last 500 chars: {final_output[-500:]}")


    print("\n" + "=" * 60)
    print("✅ ANALYSIS COMPLETE - REPORT:")
    print("=" * 60)
    print(final_output)
    print("=" * 60)

    # ==========================================
    # 🧩 STEP 3 — Parse LLM Output into JSON
    # ==========================================
    json_data = parse_report_to_json(
        final_output,
        file_path,
        criteria,
        regex_flags,           # NEW
        definition_analysis,   # NEW
        cross_ref_map          # NEW
    )
    json_data = validate_and_fix_risk_level(json_data)

    # ==========================================
    # 🔍 STEP 4 — Validate JSON
    # ==========================================
    is_valid, validation_message = validate_json_against_schema(
        json_data,
        RISK_ANALYSIS_SCHEMA
    )

    json_data['validation'] = {
        'is_valid': is_valid,
        'message': validation_message
    }
    json_data['_raw_output'] = final_output

    return json_data




def clean_markdown_artifacts(text: str) -> str:
    """Remove markdown code fences and other artifacts from text."""
    # Remove code fences
    #text = re.sub(r'```[a-zA-Z]*\n?', '', text)
    text = re.sub(r'```', '', text)
    
    # Remove extra blank lines (more than 2 consecutive)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


def calculate_risk_level(percentage: int) -> str:
    """
    Calculate risk level based on percentage and return canonical label used in reports.
    Returns one of: "LOW RISK", "MODERATE RISK", "HIGH RISK"
    Boundary logic:
      0 - 33  => LOW RISK
      34 - 66 => MODERATE RISK
      67 - 100 => HIGH RISK
    """
    try:
        p = int(round(float(percentage)))
    except Exception:
        p = 100  # fail-safe to HIGH if invalid input

    if p <= 33:
        return "LOW RISK"
    elif p <= 66:
        return "MODERATE RISK"
    else:
        return "HIGH RISK"



def validate_json_against_schema(json_data: dict, schema: dict) -> tuple:
    """Validate JSON data against the schema."""
    try:
        from jsonschema import validate, ValidationError
        validate(instance=json_data, schema=schema)
        return True, "✓ JSON data is valid according to schema"
    except ImportError:
        return True, "⚠ jsonschema library not installed. Install with: pip install jsonschema"
    except ValidationError as e:
        return False, f"✗ Validation Error: {e.message}"
    except Exception as e:
        return False, f"✗ Unexpected error during validation: {str(e)}"


def save_json_report(json_data: dict, document_path: str):
    """Save JSON report with validation and return the file path."""
    doc_path = Path(document_path)
    doc_folder = doc_path.parent
    doc_name = doc_path.stem

    # Save JSON data only
    json_filename = f"{doc_name}_risk_analysis.json"
    json_path = doc_folder / json_filename

    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(json_data, f, indent=2, ensure_ascii=False)

    print("=" * 60)
    print(f"✓ JSON data saved to: {json_path}")

    # Validate JSON against imported schema
    is_valid, message = validate_json_against_schema(json_data, RISK_ANALYSIS_SCHEMA)
    print(message)
    print("=" * 60)

    return str(json_path)


def save_json_report(json_data: dict, document_path: str):
    """Save JSON report with validation and return the file path."""
    doc_path = Path(document_path)
    doc_folder = doc_path.parent
    doc_name = doc_path.stem

    # Save JSON data only
    json_filename = f"{doc_name}_risk_analysis.json"
    json_path = doc_folder / json_filename

    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(json_data, f, indent=2, ensure_ascii=False)

    print("=" * 60)
    print(f"✓ JSON data saved to: {json_path}")

    # Validate JSON against imported schema
    is_valid, message = validate_json_against_schema(json_data, RISK_ANALYSIS_SCHEMA)
    print(message)
    print("=" * 60)

    return str(json_path)


def parse_report_to_json(text_report: str, document_path: str, criteria: list, 
                         regex_flags: dict = None, definition_analysis: dict = None, 
                         cross_ref_map: dict = None) -> dict:
    """
    FIXED VERSION - Extracts ALL hidden risks reliably
    """
    import re
    from datetime import datetime

    # NOTE: The logger and clean_markdown_artifacts are assumed to be defined elsewhere in the execution environment.
    # We must keep them in the code as requested.
    def clean_markdown_artifacts(text):
        # Placeholder implementation for clean_markdown_artifacts
        return text.replace("**", "").replace("*", "").strip()

    class Logger:
        def info(self, msg):
            # print(f"INFO: {msg}")
            pass
        def warning(self, msg):
            # print(f"WARNING: {msg}")
            pass
        def debug(self, msg):
            # print(f"DEBUG: {msg}")
            pass
        def error(self, msg):
            # print(f"ERROR: {msg}")
            pass
    
    logger = Logger()
    # End of placeholders

    text_report = text_report.replace("```", "")
    lines = text_report.split("\n")

    # ============================
    # SUMMARY EXTRACTION
    # ============================
    summary = ""
    for i, line in enumerate(lines):
        if line.strip() and not line.startswith("LEGAL DOCUMENT") and not line.startswith("==="):
            if "VENDOR" in line or "JURISDICTION" in line or "PROTECTIONS FOUND" in line:
                break
            summary = line.strip()
    summary = clean_markdown_artifacts(summary)

    hidden_risks = []
    in_hidden_section = False
    current_trap = None
    trap_number = 0

    
    END_MARKERS = [
        "DETECTION SUMMARY:",
        "VENDOR & JURISDICTION",
        "INDIAN CONTRACT ACT",
        "10XDS COMPANY",
        "UNIVERSAL NDA",
        "COUNTER-PROPOSALS",
        "RECOMMENDED COUNTER"
    ]

    logger.info("🔍 Starting hidden risks extraction...")

    for i, raw_line in enumerate(lines):
        stripped = raw_line.strip()
        if not stripped:
            continue
        upper = stripped.upper()

        # ===== START HIDDEN SECTION =====
        if not in_hidden_section:
            if "HIDDEN" in upper and ("RISK" in upper or "TRAP" in upper):
                in_hidden_section = True
                logger.info(f"  ✅ Line {i}: Hidden Risks section started")
                
                # Check if trap starts on same line as header
                trap_match = re.search(r'HIDDEN TRAP #(\d+)[:\s]+(.+)', stripped, re.IGNORECASE)
                if trap_match:
                    trap_number = int(trap_match.group(1))
                    trap_name = trap_match.group(2).strip()
                    trap_name = re.sub(r'^Name:\s*', '', trap_name, flags=re.IGNORECASE)
                    current_trap = {
                        'name': trap_name,
                        'primary_clause': '',
                        'clause_number': '',
                        'real_meaning': '',
                        'severity': ''
                    }
                    logger.info(f"    🎭 Found trap on header line: {trap_name}")
                continue

        # ===== END HIDDEN SECTION =====
        if in_hidden_section:
            # ✅ SIMPLE END CHECK (matches applocal.py)
            if any(marker in upper for marker in END_MARKERS):
                if current_trap and current_trap.get('name'):
                    hidden_risks.append(current_trap)
                    logger.info(f"  ✅ Saved final trap: {current_trap['name']}")
                logger.info(f"  🛑 Line {i}: Hidden Risks section ended")
                break  # Exit cleanly

        # ===== PARSE TRAP CONTENT =====
        if in_hidden_section:
            
            # Check for new trap starting
            trap_match = re.search(
                r'[\*\*]*\s*HIDDEN TRAP\s*#?(\d+)[:\s]*(.+)',
                stripped,
                re.IGNORECASE
            )
            
            if trap_match:
                # Save previous trap
                if current_trap and current_trap.get('name'):
                    hidden_risks.append(current_trap)
                    logger.info(f"  ✅ Saved trap #{trap_number}: {current_trap['name']}")
                
                trap_number = int(trap_match.group(1))
                trap_name = trap_match.group(2).strip()
                trap_name = re.sub(r'^Name:\s*', '', trap_name, flags=re.IGNORECASE)
                trap_name = re.sub(r'^\*+\s*', '', trap_name)
                
                current_trap = {
                    'name': trap_name,
                    'primary_clause': '',
                    'clause_number': '',
                    'real_meaning': '',
                    'severity': ''
                }
                logger.info(f"    🎭 New trap #{trap_number}: {trap_name}")
                continue

            # Parse trap fields
            if current_trap:
                
                # Name field (if separate line)
                if stripped.startswith("Name:"):
                    name = stripped.replace("Name:", "", 1).strip()
                    if name and not current_trap.get('name'):
                        current_trap['name'] = name
                        logger.info(f"      📝 Name: {name}")
                    continue

                # Primary Clause
                if re.match(r'^Primary Clause', stripped, re.IGNORECASE):
                    content = re.sub(r'^Primary Clause[:\s]*', '', stripped, flags=re.IGNORECASE).strip()
                    current_trap['primary_clause'] = content
                    
                    clause_match = re.search(r'Clause\s+(\d+[A-Za-z]?)', content, re.IGNORECASE)
                    if clause_match:
                        current_trap['clause_number'] = clause_match.group(1)
                    
                    logger.info(f"      📄 Primary Clause: {content[:50]}...")
                    continue

                # Real Meaning / Real Impact
                if re.match(r'^Real\s+(Meaning|Impact)', stripped, re.IGNORECASE):
                    content = re.sub(r'^Real\s+(Meaning|Impact)[:\s]*', '', stripped, flags=re.IGNORECASE).strip()
                    current_trap['real_meaning'] = content
                    logger.info(f"      💡 Real Meaning: {content[:50]}...")
                    continue

                # Severity
                if stripped.startswith("Severity:"):
                    severity_raw = stripped.replace("Severity:", "", 1).strip()
                    severity_clean = re.sub(r'\s*\(.*?\)\s*', '', severity_raw).strip()
                    current_trap['severity'] = severity_clean
                    logger.info(f"      ⚠️ Severity: {severity_clean}")
                    continue

    # Save last trap if exists
    if current_trap and current_trap.get('name'):
        hidden_risks.append(current_trap)
        logger.info(f"  ✅ Saved last trap: {current_trap['name']}")

    # ===== DEDUPLICATION =====
    seen_clauses = set()
    deduplicated_risks = []
    for risk in hidden_risks:
        clause_key = (risk.get('clause_number') or risk.get('primary_clause', '')[:50]).strip().lower()
        if clause_key and clause_key not in seen_clauses:
            seen_clauses.add(clause_key)
            deduplicated_risks.append(risk)
        else:
            logger.info(f"  🗑️ Duplicate removed: {risk.get('name', '<unnamed>')}")

    hidden_risks = deduplicated_risks

    logger.info(f"🎯 FINAL: Extracted {len(hidden_risks)} hidden risks")
    for idx, risk in enumerate(hidden_risks, 1):
        logger.info(f"  {idx}. {risk.get('name', 'Unknown')} - {risk.get('severity', 'N/A')}")



    # ============================
    # JURISDICTION INTELLIGENCE
    # ============================
    jurisdiction_data = {}
    in_jurisdiction = False

    for line in lines:
        if "VENDOR" in line.upper() and "JURISDICTION" in line.upper():
            in_jurisdiction = True
            continue
        if in_jurisdiction and ("INDIAN CONTRACT ACT" in line or "10XDS COMPANY" in line):
            in_jurisdiction = False
            continue
        if in_jurisdiction and ":" in line:
            key, val = line.split(":", 1)
            key_clean = key.strip().lower().replace(" ", "_")
            jurisdiction_data[key_clean] = val.strip()

        # ============================
    # INDIA LAW COMPLIANCE
    # ============================
    indian_law = {"compliant_items": [], "violations": [], "risks": []}
    in_indian_section = False
    current_section = None

    for i, line in enumerate(lines):
        stripped = line.strip()
        upper = stripped.upper()

        # Start Indian Law section
        if "INDIAN CONTRACT ACT" in upper and "COMPLIANCE" in upper:
            in_indian_section = True
            logger.info(f"📋 Line {i}: Started Indian Law section")
            continue

        # End Indian Law section
        if in_indian_section and any(marker in upper for marker in ["10XDS COMPANY", "UNIVERSAL NDA", "OVERALL RISK"]):
            logger.info(f"✅ Line {i}: Ended Indian Law section")
            logger.info(
                f"   Parsed: {len(indian_law['compliant_items'])} compliant, "
                f"{len(indian_law['violations'])} violations, "
                f"{len(indian_law['risks'])} risks"
            )
            break

        if in_indian_section and stripped:
            # Detect subsection headers
            if "COMPLIANT ITEMS" in upper or stripped.startswith("✓"):
                current_section = "compliant_items"
                logger.info(f"  ✓ Line {i}: Switched to compliant_items")
                continue

            if "VIOLATIONS FOUND" in upper or stripped.startswith("✗"):
                current_section = "violations"
                logger.info(f"  ✗ Line {i}: Switched to violations")
                continue

            if "RISKS IDENTIFIED" in upper or stripped.startswith("⚠"):
                current_section = "risks"
                logger.info(f"  ⚠ Line {i}: Switched to risks")
                continue

            # Skip "None" entries
            if stripped.lower() == "none":
                continue

            # Add content to current section
            if current_section and len(stripped) > 5:
                cleaned = stripped
                cleaned = re.sub(r'^[✓✗⚠\s\-•]+', '', cleaned).strip()

                if cleaned:
                    indian_law[current_section].append(cleaned)
                    logger.info(f"    → Added: {cleaned[:60]}...")

    # ============================
    # COMPANY POLICY - FIXED
    # ============================
    company_policy = {
        "blocking_violations": [],
        "missing_protections": [],
        "preference_gaps": []
    }
    in_company_section = False
    current_section = None

    for i, line in enumerate(lines):
        stripped = line.strip()
        upper = stripped.upper()

        # Start Company Policy section
        if "COMPANY POLICY COMPLIANCE" in upper or ("10XDS" in upper and "COMPLIANCE" in upper):
            in_company_section = True
            logger.info(f"📋 Line {i}: Started Company Policy section")
            continue

        # End Company Policy section
        if in_company_section and any(marker in upper for marker in ["UNIVERSAL NDA", "OVERALL RISK", "RECOMMENDATION"]):
            logger.info(f"✅ Line {i}: Ended Company Policy section")
            logger.info(
                f"   Parsed: {len(company_policy['blocking_violations'])} blocking, "
                f"{len(company_policy['missing_protections'])} missing, "
                f"{len(company_policy['preference_gaps'])} preferences"
            )
            break

        if in_company_section and stripped:
            # Detect subsection headers
            if "BLOCKING VIOLATIONS" in upper or stripped.startswith("🚫"):
                current_section = "blocking_violations"
                logger.info(f"  🚫 Line {i}: Switched to blocking_violations")
                continue

            if "MISSING" in upper and "PROTECTIONS" in upper:
                current_section = "missing_protections"
                logger.info(f"  ✗ Line {i}: Switched to missing_protections")
                continue

            if "PREFERENCE GAPS" in upper or stripped.startswith("ℹ"):
                current_section = "preference_gaps"
                logger.info(f"  ℹ Line {i}: Switched to preference_gaps")
                continue

            # Skip "None" entries
            if stripped.lower() == "none":
                continue

            # Add content to current section
            if current_section and len(stripped) > 5:
                cleaned = stripped
                cleaned = re.sub(r'^[🚫✗ℹ\s\-•]+', '', cleaned).strip()

                if cleaned:
                    company_policy[current_section].append(cleaned)
                    logger.info(f"    → Added: {cleaned[:60]}...")



        # ============================
    # PROTECTIONS FOUND + MISSING - IMPROVED PARSING
    # ============================
    protections_found = []
    protections_missing = []

    # First, extract the counts from the report
    found_count = 0
    missing_count = 0
    total_count = 10  # Default to 10 universal criteria

    for line in lines:
        if "Protections Found:" in line and "out of" in line:
            match = re.search(r'(\d+)\s+out of\s+(\d+)', line)
            if match:
                found_count = int(match.group(1))
                total_count = int(match.group(2))
                missing_count = total_count - found_count
                logger.info(f"📊 Protections: {found_count} found, {missing_count} missing, {total_count} total")
                break

    # Now parse the actual protection details
    in_found = False
    in_missing = False
    current_item = None

    for i, line in enumerate(lines):
        stripped = line.strip()
        upper = stripped.upper()

        # === START PROTECTIONS FOUND SECTION ===
        if "PROTECTIONS FOUND" in upper or (in_found and stripped.startswith("FOUND:")):
            if not in_found:
                in_found = True
                in_missing = False
                logger.info(f"✓ Line {i}: Started PROTECTIONS FOUND section")
            continue

        # === START PROTECTIONS MISSING SECTION ===
        if "PROTECTIONS MISSING" in upper or "MISSING:" in upper or "NOT FOUND:" in upper:
            if current_item and in_found:
                protections_found.append(current_item)
                logger.info(f"  Saved found protection: {current_item['name'][:50]}")
            in_found = False
            in_missing = True
            current_item = None
            logger.info(f"✓ Line {i}: Started PROTECTIONS MISSING section")
            continue

        # === END PROTECTIONS SECTION ===
        if in_missing and any(end in upper for end in ["RECOMMENDED COUNTER", "OVERALL RISK", "RECOMMENDATION:"]):
            if current_item:
                protections_missing.append(current_item)
                logger.info(f"  Saved missing protection: {current_item['name'][:50]}")
            logger.info(f"✓ Line {i}: Ended protections section")
            break

        # === PARSE FOUND PROTECTIONS ===
        if in_found and stripped:
            # New protection item (starts with number, checkmark, or "FOUND:")
            if (re.match(r'^\d+\.', stripped) or 
                stripped.startswith("✓") or 
                stripped.startswith("FOUND:")):

                if current_item:
                    protections_found.append(current_item)
                    logger.info(f"  Saved found protection: {current_item['name'][:50]}")

                # Extract name
                name = re.sub(r'^[\d\.\s✓]+|^FOUND:\s*', "", stripped)
                current_item = {"name": name, "clause": None, "evidence": None}

            # Clause line
            elif current_item and "Clause:" in stripped:
                current_item["clause"] = stripped.split("Clause:", 1)[1].strip()

            # Evidence line
            elif current_item and "Evidence:" in stripped:
                current_item["evidence"] = stripped.split("Evidence:", 1)[1].strip()

        # === PARSE MISSING PROTECTIONS ===
        if in_missing and stripped:
            # New missing item
            if (re.match(r'^\d+\.', stripped) or 
                stripped.startswith("✗") or 
                stripped.startswith("NOT FOUND:")):

                if current_item:
                    protections_missing.append(current_item)
                    logger.info(f"  Saved missing protection: {current_item['name'][:50]}")

                name = re.sub(r'^[\d\.\s✗]+|^NOT FOUND:\s*', "", stripped)
                current_item = {"name": name, "risk": None}

            # Risk line
            elif current_item and "Risk:" in stripped:
                current_item["risk"] = stripped.split("Risk:", 1)[1].strip()

    # === FALLBACK: Use counts if parsing failed ===
    if len(protections_found) == 0 and found_count > 0:
        logger.warning(f"⚠️ Parsing failed - using count-based fallback for {found_count} found protections")
        for i in range(found_count):
            protections_found.append({
                "name": f"Protection {i+1} (detailed parsing failed - see report)",
                "clause": None,
                "evidence": None
            })

    if len(protections_missing) == 0 and missing_count > 0:
        logger.warning(f"⚠️ Parsing failed - using count-based fallback for {missing_count} missing protections")
        for i in range(missing_count):
            protections_missing.append({
                "name": f"Missing protection {i+1} (detailed parsing failed - see report)",
                "risk": "See detailed compliance section in report"
            })

    logger.info(f"✅ Final: {len(protections_found)} found, {len(protections_missing)} missing")

         # ============================
    # COUNTER PROPOSALS - FIXED PARSING
    # ============================
    counter_proposals = []

    # Method 1: Parse strict marker format (---PROPOSAL START/END---)
    full_text = '\n'.join(lines)
    proposal_blocks = re.findall(
        r'---PROPOSAL START---\s*(.*?)\s*---PROPOSAL END---',
        full_text,
        re.DOTALL | re.IGNORECASE
    )

    logger.info(f"🔍 Found {len(proposal_blocks)} proposals with markers")

    for idx, block in enumerate(proposal_blocks):
        try:
            proposal = {}
            # Parse each field
            name_match = re.search(r'Name:\s*(.+)', block, re.IGNORECASE)
            priority_match = re.search(r'Priority:\s*(BLOCKING|HIGH|MEDIUM|LOW)', block, re.IGNORECASE)
            issue_match = re.search(r'Issue:\s*(.+?)(?=Clause:|$)', block, re.DOTALL | re.IGNORECASE)
            clause_match = re.search(r'Clause:\s*(.+?)(?=Benefit:|$)', block, re.DOTALL | re.IGNORECASE)
            benefit_match = re.search(r'Benefit:\s*(.+)', block, re.DOTALL | re.IGNORECASE)
            
            if name_match and priority_match:
                proposal['name'] = name_match.group(1).strip()
                proposal['priority'] = priority_match.group(1).upper()
                proposal['current_issue'] = issue_match.group(1).strip() if issue_match else "Issue not specified"
                proposal['suggested_clause'] = clause_match.group(1).strip() if clause_match else "Clause not specified"
                proposal['benefit'] = benefit_match.group(1).strip() if benefit_match else "Benefit not specified"
                
                counter_proposals.append(proposal)
                logger.info(f"✅ Parsed proposal {idx+1}: {proposal['name'][:50]}")
        except Exception as e:
            logger.warning(f"⚠️ Failed to parse proposal block {idx+1}: {e}")

    # Method 2: Fallback to list-based parsing if markers failed
    if len(counter_proposals) == 0:
        logger.warning("⚠️ Marker-based parsing failed, trying fallback method...")
        
        in_counter = False
        for line in lines:
            txt = line.strip()
            
            # Start section
            if "RECOMMENDED COUNTER-PROPOSALS" in txt.upper():
                in_counter = True
                continue
            
            # End section
            if txt.startswith("===") and in_counter:
                break
            
            if in_counter:
                # Match: "- Priority: HIGH - Name: Description"
                match = re.match(r'[-*]?\s*Priority:\s*(BLOCKING|HIGH|MEDIUM|LOW)\s*[-–]\s*(.+?):\s*(.+)', txt, re.IGNORECASE)
                if match:
                    proposal = {
                        'name': match.group(2).strip(),
                        'priority': match.group(1).upper(),
                        'current_issue': match.group(3).strip(),
                        'suggested_clause': "Add appropriate clause language based on issue description.",
                        'benefit': "Addresses identified compliance gap."
                    }
                    counter_proposals.append(proposal)
                    logger.info(f"✅ Fallback parsed: {proposal['name'][:50]}")

    logger.info(f"📊 TOTAL COUNTER-PROPOSALS PARSED: {len(counter_proposals)}")

    # Final validation and error handling
    if len(counter_proposals) == 0:
        logger.error("❌ CRITICAL: No counter-proposals could be parsed!")
        
        # Check if section exists at all
        proposals_section_exists = any("COUNTER" in line and "PROPOSAL" in line for line in lines)
        
        if proposals_section_exists:
            # Section exists but parsing failed - provide helpful error
            counter_proposals = [{
                "name": "Counter-Proposals Parsing Error",
                "priority": "HIGH",
                "current_issue": "Counter-proposals were generated but could not be parsed. The LLM output format may not match expected structure. Check analysis_debug.log for raw output.",
                "suggested_clause": "Review the complete analysis report section 'RECOMMENDED COUNTER-PROPOSALS' for the generated recommendations. Consider regenerating the analysis or contact support.",
                "benefit": "Ensures all identified risks have corresponding mitigation strategies."
            }]
        else:
            # No section generated at all - token limit issue
            counter_proposals = [{
                "name": "Counter-Proposals Generation Failed",
                "priority": "HIGH",
                "current_issue": "The AI system did not generate counter-proposals. This may be due to: (1) Token limit reached during analysis, (2) Document too complex, or (3) LLM timeout. Check debug_llm_output.txt for details.",
                "suggested_clause": "Please try with a shorter document (under 5 pages recommended) or split analysis into multiple sessions. Contact support if issue persists.",
                "benefit": "N/A - Error condition requires resolution"
            }]
    else:
        # Validate all proposals have required fields
        for idx, cp in enumerate(counter_proposals):
            if not cp.get('suggested_clause') or cp['suggested_clause'] == "Clause not specified":
                logger.warning(f"⚠️ Proposal '{cp['name']}' missing clause text - adding placeholder")
                cp['suggested_clause'] = f"[Detailed clause language to be provided - addresses: {cp.get('current_issue', 'issue not specified')}]"
            
            if not cp.get('current_issue'):
                cp['current_issue'] = "Issue requires further analysis"
            
            if not cp.get('benefit'):
                cp['benefit'] = "Mitigates identified compliance or legal risk"

    
    

# ============================
# RISK ASSESSMENT
# ============================
    # ============================
    risk = {}
    risk_percentage = None
    risk_level = None

    for line in lines:
        if "Risk Percentage" in line:
            matches = re.findall(r'\d+', line)
            if matches:
                risk_percentage = int(matches[0])
                risk["risk_percentage"] = risk_percentage

        if "Risk Level" in line:
            if "HIGH RISK" in line.upper():
                risk_level = "HIGH RISK"
            elif "MODERATE RISK" in line.upper():
                risk_level = "MODERATE RISK"
            elif "LOW RISK" in line.upper():
                risk_level = "LOW RISK"
            else:
                # Fallback: Calculate from percentage
                if risk_percentage is not None:
                    if risk_percentage >= 67:
                        risk_level = "HIGH RISK"
                    elif risk_percentage >= 34:
                        risk_level = "MODERATE RISK"
                    else:
                        risk_level = "LOW RISK"
                else:
                    risk_level = "MODERATE RISK"  # Default if percentage missing
            risk["risk_level"] = risk_level

    

    # ============================
    # RECOMMENDATION
    # ============================
    recommendation = ""
    rec_flag = False

    for line in lines:
        if "RECOMMENDATION" in line:
            rec_flag = True
            continue
        if rec_flag:
            if not line.strip():
                break
            recommendation += line.strip() + " "

    # ============================
    # BUILD FINAL JSON
    # ============================
    return {
        "metadata": {
            "document_path": document_path,
            "analysis_date": datetime.now().isoformat(),
            "criteria": criteria,
            "analysis_version": "3.0-hidden-risk-detection"
        },
        "summary": summary,
        "hidden_risks_detected": hidden_risks,  # ← This now works!
        "detection_methodology": {
            "regex_matches": regex_flags['total_flags'] if regex_flags else 0,
            "llm_confirmed": len(hidden_risks),
            "false_positives_filtered": max(0, 
                (regex_flags['total_flags'] - len(hidden_risks)) if regex_flags else 0
            ),
            "definition_traps_found": len(
                [r for r in hidden_risks if 'definition' in r.get('name', '').lower()]
            ),
            "cross_reference_clusters": (
                len(cross_ref_map['risk_clusters']) if cross_ref_map else 0
            )
        },
        "jurisdiction_intelligence": jurisdiction_data,
        "indian_law_compliance": indian_law,
        "company_policy_compliance": company_policy,
        "protections_found": protections_found,
        "protections_missing": protections_missing,
        "counter_proposals": counter_proposals,
        "risk_assessment": risk,
        "recommendation": recommendation.strip()
    }


def validate_and_fix_risk_level(json_data: dict) -> dict:
    """
    Ensures risk_level matches risk_percentage using strict thresholds.
    Fixes any LLM hallucinations or override errors.
    """
    risk_assessment = json_data.get('risk_assessment', {})
    
    # Extract percentage (try multiple field names)
    risk_percentage = (
        risk_assessment.get('risk_percentage') or 
        risk_assessment.get('risk_score') or 
        0
    )
    
    # Convert to int
    try:
        risk_percentage = int(float(risk_percentage))
    except (ValueError, TypeError):
        logger.warning(f"Invalid risk percentage: {risk_percentage}, defaulting to 50")
        risk_percentage = 50
    
    # Calculate correct risk level using strict thresholds
    if risk_percentage <= 33:
        correct_risk_level = "LOW RISK"
    elif risk_percentage <= 66:
        correct_risk_level = "MODERATE RISK"
    else:
        correct_risk_level = "HIGH RISK"
    
    # Get current risk level
    current_risk_level = risk_assessment.get('risk_level', '')
    
    # Fix if mismatch
    if current_risk_level != correct_risk_level:
        logger.warning(
            f"Risk level mismatch detected! "
            f"Percentage: {risk_percentage}% → Should be {correct_risk_level}, "
            f"but LLM returned {current_risk_level}. Correcting..."
        )
        risk_assessment['risk_level'] = correct_risk_level
        risk_assessment['_corrected'] = True
        risk_assessment['_original_level'] = current_risk_level
    
    # Ensure both fields exist
    risk_assessment['risk_percentage'] = risk_percentage
    risk_assessment['risk_level'] = correct_risk_level
    
    json_data['risk_assessment'] = risk_assessment
    
    logger.info(
        f"✅ Risk validation complete: {risk_percentage}% = {correct_risk_level}"
    )
    
    return json_data


@app.route('/')
def index():
    return render_template('index1.html')



    
@app.route('/analyze', methods=['POST'])
def analyze():
    try:
        # Check if file was uploaded
        if 'document' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['document']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type. Only PDF, DOCX, and TXT are allowed'}), 400
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Handle criteria properly
        custom_criteria_text = request.form.get('criteria', '').strip()
        
        # Load universal criteria as default
        universal_criteria = load_universal_criteria()
        
        if custom_criteria_text:
            custom_criteria = [c.strip() for c in custom_criteria_text.split('\n') if c.strip()]
            all_criteria = universal_criteria.copy()
            for custom in custom_criteria:
                if custom not in all_criteria:
                    all_criteria.append(custom)
            logger.info(f"Using {len(universal_criteria)} universal + {len(custom_criteria)} custom = {len(all_criteria)} total criteria")
        else:
            all_criteria = universal_criteria
            logger.info(f"Using {len(all_criteria)} universal criteria (no custom criteria provided)")
        
        # Normalize criteria to strings
        normalized_criteria = []
        for item in all_criteria:
            if isinstance(item, dict):
                priority = item.get("priority", "MEDIUM")
                category = item.get("category", "General")
                description = item.get("description", str(item))
                normalized_criteria.append(f"[{priority}] {category}: {description}")
            else:
                normalized_criteria.append(str(item))
        
        if not normalized_criteria:
            return jsonify({'error': 'No criteria available for analysis'}), 400
        
        # Analyze document
        # IMPORTANT: analyze_document must return both result JSON and final_output text
        # Adjust analyze_document accordingly:
        #   result, final_output = analyze_document(filepath, normalized_criteria)
        result = analyze_document(filepath, normalized_criteria)
        final_output = (
    result.get('_raw_output')
    or result.get('_raw_report')
    or result.get('_raw')
    or result.get('analysis_text')
    or ''
)
        if not final_output:
            # Helpful debug log so you can see what's inside result when something goes wrong
            logger.error("DEBUG: result keys: " + ", ".join(list(result.keys())))

            if isinstance(result.get('summary'), str):
                final_output = (
                    result.get('summary', '')
                    + "\n\n"
                    + result.get('recommendation', '')
                )
            elif 'saved_file_path' in result:
                final_output = (
                    f"[Raw output missing — saved JSON at {result.get('saved_file_path')}]"
                )
        try:
            with open('debug_llm_output.txt', 'w', encoding='utf-8') as f:
                f.write(final_output)
            logger.info("=== RAW LLM OUTPUT SAVED TO debug_llm_output.txt ===")
        except Exception as e:
            logger.error(f"Could not save debug_llm_output.txt: {e}")

        # Check if counter-proposals section exists
        if "COUNTER" in final_output.upper() and "PROPOSAL" in final_output.upper():
            logger.info("✓ Counter-proposals section found in output")

            start_idx = final_output.upper().find("COUNTER")
            sample = final_output[start_idx:start_idx + 1000]
            logger.info(f"Sample of counter-proposals section:\n{sample}")
        else:
            logger.error("✗ No counter-proposals section found in LLM output!")
                
        # ============================================
        # VALIDATE COUNTER-PROPOSALS BEFORE SENDING
        # ============================================
        counter_proposals = result.get('counter_proposals', [])

        logger.info(f"Validation Check: {len(counter_proposals)} counter-proposals found")

        if len(counter_proposals) == 0:
            logger.error("CRITICAL: No counter-proposals generated!")

            # Check if raw output contained proposals
            if isinstance(final_output, str) and "COUNTER-PROPOSALS" in final_output.upper():
                logger.error("Proposals exist in output but parsing FAILED")
                start_idx = final_output.upper().find("COUNTER-PROPOSALS")
                sample = (
                    final_output[start_idx:start_idx + 500]
                    if start_idx != -1 else final_output[:500]
                )
                logger.error(f"Output sample around COUNTER-PROPOSALS:\n{sample}")
            else:
                logger.error("LLM did not generate counter-proposals section at all")

            # Inject fallback error proposal
            result['counter_proposals'] = [{
                "name": "Counter-Proposals Generation Failed",
                "priority": "HIGH",
                "current_issue": (
                    "The AI system did not generate counter-proposals. This may be due to: "
                    "(1) Token limit reached, (2) Complex document analysis, or (3) Parsing error."
                ),
                "suggested_clause": (
                    "Please try again with a shorter document, or contact support. "
                    "Check the analysis_debug.log file for detailed error information."
                ),
                "benefit": "N/A - Error condition"
            }]

        else:
            # Validate each proposal
            valid_count = 0
            valid_proposals = []

            for idx, cp in enumerate(counter_proposals, 1):
                clause = cp.get("suggested_clause", "")

                if clause and clause != "[Clause not generated - regenerate report]":
                    valid_count += 1
                    valid_proposals.append(cp)
                    logger.info(
                        f"Proposal #{idx}: {cp.get('name', 'Unknown')[:50]} | {len(clause)} chars"
                    )
                else:
                    logger.error(
                        f"Proposal #{idx} '{cp.get('name', 'Unknown')}' has NO CLAUSE TEXT"
                    )

            logger.info(
                f"FINAL: {valid_count}/{len(counter_proposals)} proposals have valid clause text"
            )

            # Keep only valid ones if any exist
            if valid_proposals:
                result['counter_proposals'] = valid_proposals

        # -------------------------------
        # 4. Save JSON to file
        # -------------------------------
        try:
            json_file_path = save_json_report(result, filepath)
            result['saved_file_path'] = json_file_path
            logger.info(f"Analysis complete. JSON saved to: {json_file_path}")

        except Exception as e:
            logger.warning(f"Could not save JSON file: {e}")

        # -------------------------------
        # 5. Final response
        # -------------------------------
        logger.info(
            f"Sending response with {len(result.get('counter_proposals', []))} counter-proposals"
        )
        return jsonify(result)

    except Exception as e:
        import traceback
        logger.error(f"Error in analyze route: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500
    

    
@app.route('/preview-criteria', methods=['GET'])
def preview_criteria():
    """Return universal criteria for frontend preview."""
    try:
        criteria = load_universal_criteria()
        return jsonify({
            'success': True,
            'count': len(criteria),
            'criteria': criteria
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/debug-last-analysis', methods=['GET'])
def debug_last_analysis():
    """Return the last analysis debug log for troubleshooting."""
    try:
        with open('analysis_debug.log', 'r') as f:
            logs = f.read()
        return jsonify({
            'success': True,
            'logs': logs[-5000:]  # Last 5000 chars
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


    
if __name__ == '__main__':
    app.run(debug=True, host='127.0.0.1', port=5000)


